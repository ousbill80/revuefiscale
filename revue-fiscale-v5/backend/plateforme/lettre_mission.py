"""Lettre de mission — livrable Word (.docx) du cadrage.

Normes professionnelles (expertise comptable / conseil fiscal) : toute
mission est cadrée par une lettre de mission signée avant les travaux.
Ce module assemble le document depuis les données déjà saisies au wizard
(mission + contribuable + tenant) — aucune donnée inventée : tout champ
manquant est rendu « [à compléter] ».

Aucun taux, seuil ou condition fiscale ici — document d'engagement.
"""
from __future__ import annotations

import io
import re
import unicodedata
from datetime import date, datetime, timezone
from typing import Any, Final

from docx import Document
from sqlalchemy import text
from sqlalchemy.orm import Session

from backend.plateforme.contexte import contexte_tenant
from backend.plateforme.missions import LIBELLES_ENGAGEMENT

# Libellés courants des codes pivot `impot` — miroir de
# frontend/mission/src/impotLabels.ts (docs/08-glossaire.md). Pas de barème.
LIBELLES_IMPOT: Final[dict[str, str]] = {
    "BIC": "Bénéfices industriels et commerciaux",
    "TVA": "Taxe sur la valeur ajoutée",
    "RAS": "Retenue à la source (notamment non-résidents)",
    "ITS": "Impôt sur les traitements et salaires",
    "CE": "Contribution employeur",
    "IRC": "Impôt sur le revenu des créances",
    "IRVM": "Impôt sur le revenu des valeurs mobilières",
    "PAT": "Contribution des patentes",
    "FONC": "Impôt foncier",
    "ENR": "Droits d'enregistrement",
    "TIMBRE": "Droit de timbre",
    "OBL": "Obligations déclaratives (ETII, registres…)",
    "OBNL": "Organismes à but non lucratif",
    "RA": "Revue analytique (contrôles de cohérence)",
}

A_COMPLETER: Final = "[à compléter]"


class ErreurLettreMission(Exception):
    """Echec de génération de la lettre de mission (mission introuvable…)."""


def _champ(valeur: object | None) -> str:
    """Valeur affichable — jamais d'invention : manquant = « [à compléter] »."""
    if valeur is None:
        return A_COMPLETER
    texte_ = str(valeur).strip()
    return texte_ or A_COMPLETER


def _fmt_montant(montant: object) -> str:
    try:
        from decimal import Decimal

        return f"{Decimal(str(montant)):,.0f}".replace(",", " ")
    except Exception:
        return str(montant)


def nom_fichier_lettre(denomination: object | None, exercice: object | None) -> str:
    """lettre_mission_{NOM}_{exercice}.docx — nom ASCII sûr pour l'en-tête HTTP."""
    brut = str(denomination or "client")
    sans_accents = (
        unicodedata.normalize("NFKD", brut).encode("ascii", "ignore").decode("ascii")
    )
    nom = re.sub(r"[^A-Za-z0-9]+", "_", sans_accents).strip("_").upper() or "CLIENT"
    exo = str(exercice or A_COMPLETER).strip() or "exercice"
    exo = re.sub(r"[^A-Za-z0-9]+", "_", exo) or "exercice"
    return f"lettre_mission_{nom}_{exo}.docx"


def collecter_donnees_lettre(
    session: Session, tenant_id: int, mission_id: int
) -> dict[str, Any]:
    """Lecture seule (RLS via contexte_tenant) : mission + contribuable + objectifs.

    L'identité du cabinet (table tenant, sans RLS) est lue par id = tenant_id,
    même garde que /api/v1/auth/connexion.
    """
    from backend.plateforme.missions import serialiser_mission
    from backend.plateforme.objectifs import lister_objectifs_en_contexte
    from backend.plateforme.objectifs_fiscaux import (
        lister_objectifs_fiscaux_en_contexte,
    )

    with contexte_tenant(session, tenant_id):
        row = session.execute(
            text(
                "SELECT m.id, m.contribuable_id, m.exercice, m.statut, "
                "m.version_referentiel_id, m.cree_le, m.type_engagement, "
                "m.perimetre_impots, m.exclusions_declarees, m.seuil_signification, "
                "c.denomination AS contribuable_denomination, c.ncc, c.rccm, "
                "c.forme_juridique, c.regime_fiscal, c.siege_social, c.commune, "
                "c.centre_impots "
                "FROM mission m JOIN contribuable c ON c.id = m.contribuable_id "
                "WHERE m.id = :m"
            ),
            {"m": mission_id},
        ).mappings().one_or_none()
        if row is None:
            raise ErreurLettreMission(f"mission {mission_id} introuvable")
        objectifs = lister_objectifs_en_contexte(session, mission_id)
        objectifs_fiscaux = lister_objectifs_fiscaux_en_contexte(session, mission_id)

    mission = serialiser_mission(
        dict(row), objectifs=objectifs, objectifs_fiscaux=objectifs_fiscaux
    )

    cabinet = session.execute(
        text(
            "SELECT denomination, ncc, rccm, forme_juridique, siege_social, "
            "commune, centre_impots, capital_social "
            "FROM tenant WHERE id = :t"
        ),
        {"t": tenant_id},
    ).mappings().one_or_none()

    return {
        "mission": mission,
        "contribuable": {
            "denomination": row["contribuable_denomination"],
            "ncc": row["ncc"],
            "rccm": row["rccm"],
            "forme_juridique": row["forme_juridique"],
            "regime_fiscal": row["regime_fiscal"],
            "siege_social": row["siege_social"],
            "commune": row["commune"],
            "centre_impots": row["centre_impots"],
        },
        "cabinet": dict(cabinet) if cabinet is not None else {},
    }


def rendre_lettre_mission_docx(donnees: dict[str, Any]) -> bytes:
    """Assemble le .docx — structure professionnelle, prête à personnaliser."""
    mission: dict[str, Any] = donnees.get("mission") or {}
    client: dict[str, Any] = donnees.get("contribuable") or {}
    cabinet: dict[str, Any] = donnees.get("cabinet") or {}

    exercice = mission.get("exercice")
    type_eng = str(mission.get("type_engagement") or "autre")
    libelle_eng = LIBELLES_ENGAGEMENT.get(type_eng, type_eng)

    doc = Document()

    # En-tête cabinet (émetteur)
    doc.add_paragraph(_champ(cabinet.get("denomination")).upper())
    doc.add_paragraph(
        f"Forme juridique : {_champ(cabinet.get('forme_juridique'))} — "
        f"RCCM : {_champ(cabinet.get('rccm'))} — NCC : {_champ(cabinet.get('ncc'))}"
    )
    capital = cabinet.get("capital_social")
    if capital is not None:
        doc.add_paragraph(f"Capital social : {_fmt_montant(capital)} FCFA")
    doc.add_paragraph(
        f"Siège : {_champ(cabinet.get('siege_social'))} — "
        f"{_champ(cabinet.get('commune'))}"
    )
    doc.add_paragraph(
        f"Centre des impôts de rattachement : {_champ(cabinet.get('centre_impots'))}"
    )

    # Destinataire
    doc.add_paragraph("")
    doc.add_paragraph(f"À l'attention de la Direction de {_champ(client.get('denomination'))}")
    doc.add_paragraph(
        f"Siège : {_champ(client.get('siege_social'))} — {_champ(client.get('commune'))}"
    )
    doc.add_paragraph(
        f"{_champ(cabinet.get('commune'))}, le {date.today().strftime('%d/%m/%Y')}"
    )

    # Objet
    doc.add_heading("Lettre de mission", level=1)
    doc.add_paragraph(
        f"Objet : mission de revue fiscale — {libelle_eng} — "
        f"exercice {_champ(exercice)}."
    )
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph(
        "Nous vous remercions de la confiance que vous nous témoignez. "
        "La présente lettre a pour objet de définir les termes et conditions "
        "de notre intervention, conformément aux normes professionnelles "
        "applicables. Elle doit être signée par les deux parties avant le "
        "démarrage des travaux."
    )

    # 1. Contexte
    doc.add_heading("1. Contexte et identification", level=2)
    doc.add_paragraph(
        f"Entité contrôlée : {_champ(client.get('denomination'))} "
        f"({_champ(client.get('forme_juridique'))})",
        style="List Bullet",
    )
    doc.add_paragraph(f"NCC : {_champ(client.get('ncc'))}", style="List Bullet")
    doc.add_paragraph(f"RCCM : {_champ(client.get('rccm'))}", style="List Bullet")
    doc.add_paragraph(
        f"Régime fiscal déclaré : {_champ(client.get('regime_fiscal'))}",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"Centre des impôts de rattachement : {_champ(client.get('centre_impots'))}",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"Exercice contrôlé : {_champ(exercice)}", style="List Bullet"
    )
    doc.add_paragraph(
        "La mission consiste en une revue fiscale de l'exercice visé, "
        "réalisée sur la base des documents comptables et fiscaux fournis "
        "par l'entité."
    )

    # 2. Nature et étendue des travaux
    doc.add_heading("2. Nature et étendue des travaux", level=2)
    doc.add_paragraph(f"Type d'engagement : {libelle_eng}.")
    perimetre = mission.get("perimetre_impots")
    if isinstance(perimetre, (list, tuple)) and len(perimetre) > 0:
        doc.add_paragraph("Impôts et taxes inclus dans le périmètre convenu :")
        for code in perimetre:
            c = str(code).strip().upper()
            libelle = LIBELLES_IMPOT.get(c)
            doc.add_paragraph(
                f"{c} — {libelle}" if libelle else c, style="List Bullet"
            )
    else:
        doc.add_paragraph(
            "Périmètre d'impôts : l'ensemble des impôts et taxes applicables "
            "à l'entité (périmètre non restreint)."
        )
    objectifs = mission.get("objectifs") or []
    libelles_obj = [
        str(o.get("libelle") or "").strip()
        for o in objectifs
        if isinstance(o, dict) and str(o.get("libelle") or "").strip()
    ]
    if libelles_obj:
        doc.add_paragraph("Objectifs convenus avec l'entité :")
        for lib in libelles_obj:
            doc.add_paragraph(lib, style="List Bullet")
    else:
        doc.add_paragraph(f"Objectifs convenus : {A_COMPLETER}")

    # 3. Limites et exclusions
    doc.add_heading("3. Limites et exclusions", level=2)
    exclusions = mission.get("exclusions_declarees")
    if exclusions and str(exclusions).strip():
        doc.add_paragraph(f"Exclusions déclarées : {str(exclusions).strip()}")
    else:
        doc.add_paragraph(f"Exclusions déclarées : {A_COMPLETER}")
    doc.add_paragraph(
        "La présente mission ne constitue pas un audit ni une certification "
        "des comptes. Elle ne peut se substituer aux obligations déclaratives "
        "de l'entité ni garantir l'absence de redressement en cas de contrôle "
        "de l'administration fiscale."
    )

    # 4. Seuil de signification (uniquement si renseigné)
    seuil = mission.get("seuil_signification")
    if seuil is not None and str(seuil).strip() != "":
        doc.add_heading("4. Seuil de signification", level=2)
        doc.add_paragraph(
            f"Le seuil de signification convenu pour la mission est fixé à "
            f"{_fmt_montant(seuil)} FCFA. Les constats d'un montant inférieur "
            "pourront être signalés sans faire l'objet de développements "
            "détaillés."
        )

    # 5. Obligations réciproques
    doc.add_heading("5. Obligations réciproques", level=2)
    doc.add_paragraph(
        "Le client s'engage à mettre à notre disposition, dans les délais "
        "convenus, l'ensemble des documents, informations et pièces "
        "justificatives nécessaires à la mission, et garantit leur sincérité "
        "et leur exhaustivité.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Le cabinet s'engage à réaliser la mission avec diligence et "
        "conscience professionnelle, et est tenu au secret professionnel "
        "sur l'ensemble des informations portées à sa connaissance.",
        style="List Bullet",
    )

    # 6. Durée
    doc.add_heading("6. Durée de la mission", level=2)
    doc.add_paragraph(
        f"La mission porte sur l'exercice {_champ(exercice)}. Elle débute à "
        "la signature de la présente lettre et s'achève à la remise du "
        "rapport de revue fiscale. Calendrier détaillé : [à compléter]."
    )
    doc.add_paragraph(
        "Honoraires et modalités de facturation : [à compléter]."
    )

    doc.add_paragraph(
        "Nous vous prions de bien vouloir nous retourner un exemplaire de la "
        "présente lettre revêtu de votre signature, précédée de la mention "
        "« Bon pour accord »."
    )
    doc.add_paragraph(
        "Nous vous prions d'agréer, Madame, Monsieur, l'expression de nos "
        "salutations distinguées."
    )

    # Blocs signature
    doc.add_heading("Signatures", level=2)
    table = doc.add_table(rows=4, cols=2)
    gauche, droite = table.rows[0].cells
    gauche.text = f"Pour le cabinet : {_champ(cabinet.get('denomination'))}"
    droite.text = f"Pour le client : {_champ(client.get('denomination'))}"
    table.rows[1].cells[0].text = "Nom et qualité : [à compléter]"
    table.rows[1].cells[1].text = "Nom et qualité : [à compléter]"
    table.rows[2].cells[0].text = "Date : [jj/mm/aaaa]"
    table.rows[2].cells[1].text = "Date : [jj/mm/aaaa]"
    table.rows[3].cells[0].text = "Signature :"
    table.rows[3].cells[1].text = "Signature (« Bon pour accord ») :"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generer_lettre_mission(
    session: Session, tenant_id: int, mission_id: int
) -> tuple[bytes, str]:
    """Contenu .docx + nom de fichier — point d'entrée de la route."""
    donnees = collecter_donnees_lettre(session, tenant_id, mission_id)
    contenu = rendre_lettre_mission_docx(donnees)
    nom = nom_fichier_lettre(
        donnees["contribuable"].get("denomination"),
        donnees["mission"].get("exercice"),
    )
    return contenu, nom


# ═════════════════════════════════════════════════════════════════════
# Lettre de mission IMPRIMABLE (JSON) — document contractuel de cadrage
# ═════════════════════════════════════════════════════════════════════
#
# POURQUOI : au CADRAGE, avant de démarrer les travaux, le fiscaliste
# remet au client la lettre de mission imprimée depuis le navigateur
# (même mécanisme que le dossier de synthèse) : identification
# cabinet/client, exercice revu, régime et principales obligations
# déclaratives du régime (dédupliquées, SANS les dates), nature et
# limites de la mission (revue fiscale consultative), obligations
# réciproques, confidentialité, honoraires convenus et signatures.
#
# Assemblage DÉTERMINISTE et CONSULTATIF (aucun LLM) : l'identité est
# lue par le MODULE EXISTANT du dossier de synthèse
# (backend.plateforme.dossier_mission._bloc_identite) et les
# obligations proviennent de l'échéancier fiscal de la mission
# (backend.plateforme.echeancier_fiscal.echeancier_mission, qui
# délègue à construire_echeancier). Aucun calcul n'est dupliqué.
# AUCUNE écriture en base ici (journal de consultation côté route).

#: Blocs de la lettre — clés STABLES garanties par l'assembleur.
BLOCS_LETTRE: Final[tuple[str, ...]] = (
    "identite",
    "objet",
    "perimetre",
    "limites",
    "obligations_reciproques",
    "confidentialite",
    "honoraires",
    "signatures",
)

TEXTE_OBJET: Final[str] = (
    "La présente lettre de mission a pour objet de définir les "
    "conditions d'intervention du Cabinet auprès du Client pour la "
    "réalisation d'une mission de revue fiscale consultative portant "
    "sur l'exercice visé en tête de lettre. La revue consiste à "
    "examiner, sur la base des pièces et informations communiquées par "
    "le Client, le respect de ses principales obligations fiscales "
    "déclaratives et de paiement, à identifier les zones de risque et "
    "à restituer au Client des constats et des recommandations à "
    "caractère strictement consultatif."
)

TEXTE_PERIMETRE: Final[str] = (
    "Le périmètre de la revue couvre, pour le régime d'imposition du "
    "Client, les principales obligations déclaratives et de paiement "
    "rappelées ci-après. Cette liste, établie d'après la pratique "
    "déclarative usuelle en Côte d'Ivoire, est indicative : le "
    "calendrier officiel de la Direction générale des impôts prévaut "
    "en toutes circonstances."
)

TEXTE_LIMITES: Final[str] = (
    "La mission est exclusivement consultative. Elle ne constitue pas "
    "un audit ni une certification des comptes ou des déclarations du "
    "Client, ni une garantie contre un contrôle ou un redressement de "
    "l'administration fiscale. Les constats et recommandations sont "
    "émis au vu des seules pièces et informations communiquées par le "
    "Client, sans vérification exhaustive de leur exactitude. Le "
    "Cabinet ne se substitue ni au Client, qui demeure seul "
    "responsable de ses déclarations et de ses paiements, ni à "
    "l'administration fiscale, seule habilitée à prendre position. Le "
    "Client reste seul décideur des suites à donner aux "
    "recommandations."
)

TEXTE_OBLIGATIONS_RECIPROQUES: Final[str] = (
    "Le Cabinet s'engage à exécuter la mission avec diligence et "
    "conformément aux règles de l'art, à affecter à la mission des "
    "intervenants disposant des compétences requises, à informer le "
    "Client de tout point significatif relevé au cours des travaux et "
    "à restituer ses conclusions dans les délais convenus. Le Client "
    "s'engage à communiquer au Cabinet, de manière exhaustive et "
    "sincère, l'ensemble des documents et informations nécessaires à "
    "la revue, à en garantir la sincérité et à informer le Cabinet "
    "sans délai de tout événement susceptible d'affecter le "
    "déroulement de la mission."
)

TEXTE_CONFIDENTIALITE: Final[str] = (
    "Le Cabinet est tenu au secret professionnel. Les documents et "
    "informations communiqués par le Client sont traités de manière "
    "confidentielle, ne sont utilisés que pour les besoins de la "
    "mission et ne sont communiqués à aucun tiers sans l'accord "
    "préalable et écrit du Client, sauf obligation légale ou "
    "réglementaire. Cette obligation de confidentialité survit au "
    "terme de la mission."
)

TEXTE_HONORAIRES_CONVENUS: Final[str] = (
    "Les honoraires convenus entre les parties pour la présente "
    "mission s'élèvent au montant indiqué ci-après, hors taxes et "
    "hors débours éventuels, payables selon les modalités arrêtées "
    "d'un commun accord."
)

TEXTE_HONORAIRES_A_CONVENIR: Final[str] = (
    "Les honoraires de la présente mission seront arrêtés d'un commun "
    "accord entre les parties et feront l'objet d'une facturation du "
    "Cabinet, hors taxes et hors débours éventuels."
)

MENTION_SIGNATURE_CABINET: Final[str] = "Pour le Cabinet"
MENTION_SIGNATURE_CLIENT: Final[str] = "Pour le Client"
MENTION_LU_APPROUVE: Final[str] = (
    "Signature précédée de la mention manuscrite « Lu et approuvé »."
)

MENTION_NOTE_LETTRE: Final[str] = (
    "Modèle indicatif de lettre de mission, assemblé de façon "
    "déterministe à partir des informations de cadrage saisies dans "
    "l'application — à relire et à adapter avant remise au client : "
    "le cabinet reste responsable de sa lettre de mission."
)


class ErreurLettreIntrouvable(ErreurLettreMission):
    """Mission hors périmètre du tenant — 404 côté route."""


def _deduire_obligations(
    echeances: list[dict[str, Any]] | None,
) -> list[dict[str, str]]:
    """Couples (impôt, obligation) UNIQUES, sans dates, ordre d'apparition.

    L'échéancier répète chaque obligation à chaque échéance (12 lignes
    de TVA mensuelle…) : la lettre ne cite chaque obligation qu'une
    seule fois. L'ordre suit la première occurrence (échéancier déjà
    trié par date limite) — stable et déterministe.
    """
    vus: set[tuple[str, str]] = set()
    obligations: list[dict[str, str]] = []
    for e in echeances or []:
        cle = (str(e.get("impot") or ""), str(e.get("obligation") or ""))
        if cle in vus or cle == ("", ""):
            continue
        vus.add(cle)
        obligations.append({"impot": cle[0], "obligation": cle[1]})
    return obligations


def assembler_lettre(
    identite: dict[str, Any],
    obligations: list[dict[str, Any]] | None,
    genere_le: str | None = None,
) -> dict[str, Any]:
    """PUR — assemble la lettre de mission (clés stables, testable).

    ``identite`` : bloc d'identité du dossier de synthèse (cabinet,
    contribuable, exercice, régime, honoraires en str Decimal ou None) ;
    ``obligations`` : items d'échéancier, avec ou sans doublons —
    dédupliqués ici par couple (impôt, obligation), dates IGNORÉES ;
    ``genere_le`` : horodatage ISO UTC (fourni pour les tests, sinon
    maintenant). Toutes les clés de :data:`BLOCS_LETTRE` sont toujours
    présentes, plus ``genere_le`` et ``note`` (modèle indicatif).
    """
    honoraires = identite.get("honoraires")
    montant = str(honoraires) if honoraires is not None else None
    return {
        "identite": dict(identite),
        "objet": TEXTE_OBJET,
        "perimetre": {
            "texte": TEXTE_PERIMETRE,
            "regime": identite.get("regime"),
            "obligations": _deduire_obligations(obligations),
        },
        "limites": TEXTE_LIMITES,
        "obligations_reciproques": TEXTE_OBLIGATIONS_RECIPROQUES,
        "confidentialite": TEXTE_CONFIDENTIALITE,
        "honoraires": {
            "montant": montant,
            "texte": (
                TEXTE_HONORAIRES_CONVENUS
                if montant is not None
                else TEXTE_HONORAIRES_A_CONVENIR
            ),
        },
        "signatures": {
            "cabinet": {
                "titre": MENTION_SIGNATURE_CABINET,
                "denomination": str(identite.get("cabinet") or ""),
            },
            "client": {
                "titre": MENTION_SIGNATURE_CLIENT,
                "denomination": str(identite.get("contribuable") or ""),
            },
            "mention": MENTION_LU_APPROUVE,
        },
        "genere_le": genere_le
        or datetime.now(timezone.utc).replace(microsecond=0).isoformat(),
        "note": MENTION_NOTE_LETTRE,
    }


def lettre_mission(
    session: Session, tenant_id: int, mission_id: int
) -> dict[str, Any]:
    """Lettre de mission imprimable (LECTURE SEULE, RLS).

    RÉUTILISE la lecture d'identité du dossier de synthèse
    (:func:`backend.plateforme.dossier_mission._bloc_identite`) et
    l'échéancier fiscal de la mission
    (:func:`backend.plateforme.echeancier_fiscal.echeancier_mission`,
    qui délègue à ``construire_echeancier`` — régime du profil JSON de
    la mission, DGE détectée depuis le centre des impôts). Mission hors
    tenant → :class:`ErreurLettreIntrouvable` (404 côté route).
    AUCUNE écriture en base (journal de consultation côté route).
    """
    from backend.plateforme.dossier_mission import (
        ErreurDossierIntrouvable,
        _bloc_identite,
    )
    from backend.plateforme.echeancier_fiscal import echeancier_mission

    try:
        identite = _bloc_identite(session, tenant_id, mission_id)
    except ErreurDossierIntrouvable as e:
        raise ErreurLettreIntrouvable(str(e)) from e
    echeancier = echeancier_mission(session, tenant_id, mission_id)
    return assembler_lettre(identite, echeancier.get("echeances"))
