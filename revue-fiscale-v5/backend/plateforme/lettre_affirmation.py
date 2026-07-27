"""Lettre d'affirmation de la direction — document Word (.docx) à signer.

Pratique normée (esprit NEP 580 / ISA 580 adaptée à la revue fiscale) :
avant la clôture de la mission, le cabinet fait signer au dirigeant une
lettre par laquelle la direction confirme l'exhaustivité des informations
transmises (déclarations, balances/FEC, litiges et contrôles en cours,
exhaustivité des passifs fiscaux, réponses données aux demandes).

Particularité : le document est à en-tête du CLIENT (expéditeur =
contribuable) et adressé au CABINET — pré-rempli et daté, à faire signer
par le représentant légal.

Assemblage DÉTERMINISTE (aucun appel LLM, lecture seule sous RLS via
``contexte_tenant``). Si la mission a des risques ouverts ou des
conclusions en anomalie (dernière exécution), la lettre mentionne leur
NOMBRE — la direction affirme alors n'avoir connaissance d'aucun autre
élément. Aucun taux ni seuil fiscal ici — document d'engagement.
"""
from __future__ import annotations

import io
import re
import unicodedata
from datetime import date
from typing import Any, Final

from docx import Document
from sqlalchemy import text
from sqlalchemy.orm import Session

from backend.plateforme.contexte import contexte_tenant
from backend.plateforme.demande_renseignements import A_COMPLETER, _champ

# Statuts de risque considérés comme non traités — mêmes valeurs que
# _STATUTS_RISQUE_OUVERTS de controle_cloture.py (STATUTS_NON_CLOS).
STATUTS_RISQUE_OUVERTS: Final[tuple[str, ...]] = ("ouvert", "en_traitement")


class ErreurLettreAffirmation(Exception):
    """Echec de génération de la lettre d'affirmation."""


class ErreurLettreAffirmationIntrouvable(ErreurLettreAffirmation):
    """Mission hors périmètre du tenant — 404 côté route."""


def nom_fichier_lettre_affirmation(
    denomination: object | None, exercice: object | None
) -> str:
    """lettre_affirmation_{NOM}_{exercice}.docx — ASCII sûr (HTTP)."""
    brut = str(denomination or "client")
    sans_accents = (
        unicodedata.normalize("NFKD", brut).encode("ascii", "ignore").decode("ascii")
    )
    nom = re.sub(r"[^A-Za-z0-9]+", "_", sans_accents).strip("_").upper() or "CLIENT"
    exo = str(exercice or A_COMPLETER).strip() or "exercice"
    exo = re.sub(r"[^A-Za-z0-9]+", "_", exo) or "exercice"
    return f"lettre_affirmation_{nom}_{exo}.docx"


def _nb_anomalies_derniere_execution(session: Session, mission_id: int) -> int:
    """Conclusions en anomalie de la DERNIÈRE exécution — 0 si aucune.

    Mêmes tables execution/conclusion que ``courrier_envoi_rapport``.
    """
    execution_id = session.execute(
        text(
            "SELECT id FROM execution "
            "WHERE mission_id = :m ORDER BY id DESC LIMIT 1"
        ),
        {"m": mission_id},
    ).scalar_one_or_none()
    if execution_id is None:
        return 0
    return int(
        session.execute(
            text(
                "SELECT count(*) FROM conclusion "
                "WHERE execution_id = :e AND statut = 'anomalie'"
            ),
            {"e": int(execution_id)},
        ).scalar_one()
    )


def _nb_risques_ouverts(session: Session, contribuable_id: int) -> int:
    """Risques du contribuable encore ouverts ou en traitement."""
    return int(
        session.execute(
            text(
                "SELECT count(*) FROM risque "
                "WHERE contribuable_id = :c AND statut = ANY(:ouverts)"
            ),
            {"c": contribuable_id, "ouverts": list(STATUTS_RISQUE_OUVERTS)},
        ).scalar_one()
    )


def collecter_donnees_lettre_affirmation(
    session: Session, tenant_id: int, mission_id: int
) -> dict[str, Any]:
    """Lecture seule (RLS via contexte_tenant) : mission + client + compteurs.

    Mission hors tenant → :class:`ErreurLettreAffirmationIntrouvable` (404).
    L'identité du cabinet (table tenant, sans RLS) est lue par
    id = tenant_id — même garde que /api/v1/auth/connexion.
    """
    with contexte_tenant(session, tenant_id):
        row = session.execute(
            text(
                "SELECT m.id, m.exercice, m.contribuable_id, "
                "c.denomination AS contribuable_denomination, c.ncc, "
                "c.forme_juridique, c.siege_social, c.commune "
                "FROM mission m JOIN contribuable c ON c.id = m.contribuable_id "
                "WHERE m.id = :m"
            ),
            {"m": mission_id},
        ).mappings().one_or_none()
        if row is None:
            raise ErreurLettreAffirmationIntrouvable(
                f"mission {mission_id} introuvable"
            )
        nb_anomalies = _nb_anomalies_derniere_execution(session, mission_id)
        nb_risques_ouverts = _nb_risques_ouverts(
            session, int(row["contribuable_id"])
        )

    cabinet = session.execute(
        text(
            "SELECT denomination, siege_social, commune "
            "FROM tenant WHERE id = :t"
        ),
        {"t": tenant_id},
    ).mappings().one_or_none()

    return {
        "mission": {
            "id": int(row["id"]),
            "exercice": row["exercice"],
        },
        "contribuable": {
            "denomination": row["contribuable_denomination"],
            "ncc": row["ncc"],
            "forme_juridique": row["forme_juridique"],
            "siege_social": row["siege_social"],
            "commune": row["commune"],
        },
        "cabinet": dict(cabinet) if cabinet is not None else {},
        "nb_anomalies": nb_anomalies,
        "nb_risques_ouverts": nb_risques_ouverts,
    }


def rendre_lettre_affirmation_docx(donnees: dict[str, Any]) -> bytes:
    """Assemble le .docx — en-tête CLIENT, affirmations, signature dirigeant."""
    mission: dict[str, Any] = donnees.get("mission") or {}
    client: dict[str, Any] = donnees.get("contribuable") or {}
    cabinet: dict[str, Any] = donnees.get("cabinet") or {}
    nb_anomalies = int(donnees.get("nb_anomalies") or 0)
    nb_risques_ouverts = int(donnees.get("nb_risques_ouverts") or 0)

    exercice = mission.get("exercice")
    doc = Document()

    # En-tête CLIENT (expéditeur) — la lettre émane de la direction.
    doc.add_paragraph(_champ(client.get("denomination")).upper())
    doc.add_paragraph(
        f"Forme juridique : {_champ(client.get('forme_juridique'))}"
    )
    ncc = str(client.get("ncc") or "").strip()
    if ncc:
        doc.add_paragraph(f"NCC : {ncc}")
    doc.add_paragraph(
        f"Siège : {_champ(client.get('siege_social'))} — "
        f"{_champ(client.get('commune'))}"
    )

    # Destinataire : le cabinet.
    doc.add_paragraph("")
    doc.add_paragraph(
        f"À l'attention de : {_champ(cabinet.get('denomination'))}"
    )
    doc.add_paragraph(
        f"Siège : {_champ(cabinet.get('siege_social'))} — "
        f"{_champ(cabinet.get('commune'))}"
    )
    doc.add_paragraph(
        f"{_champ(client.get('commune'))}, le {date.today().strftime('%d/%m/%Y')}"
    )

    # Objet
    doc.add_heading("Lettre d'affirmation de la direction", level=1)
    doc.add_paragraph(
        f"Objet : Lettre d'affirmation — revue fiscale exercice "
        f"{_champ(exercice)}."
    )
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph(
        "Dans le cadre de votre mission de revue fiscale portant sur "
        f"l'exercice {_champ(exercice)} (mission n° {mission.get('id')}), et "
        "en notre qualité de représentant légal de "
        f"{_champ(client.get('denomination'))}, nous vous confirmons, au "
        "mieux de notre connaissance et en toute bonne foi, les affirmations "
        "suivantes :"
    )

    # Affirmations standard.
    doc.add_heading("Affirmations de la direction", level=2)
    doc.add_paragraph(
        "La comptabilité de l'entité, les balances et le fichier des "
        "écritures comptables (FEC) qui vous ont été remis sont exhaustifs, "
        "sincères et conformes aux livres et registres de l'entité.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "L'ensemble des déclarations fiscales souscrites au titre de "
        f"l'exercice {_champ(exercice)} vous a été communiqué, sans omission.",
        style="List Bullet",
    )
    if nb_risques_ouverts > 0 or nb_anomalies > 0:
        precisions: list[str] = []
        if nb_risques_ouverts > 0:
            precisions.append(
                f"{nb_risques_ouverts} risque(s) fiscal(aux) encore ouvert(s)"
            )
        if nb_anomalies > 0:
            precisions.append(
                f"{nb_anomalies} conclusion(s) en anomalie relevée(s) lors "
                "de la dernière exécution de vos contrôles"
            )
        doc.add_paragraph(
            "En dehors des éléments déjà portés à votre connaissance dans "
            "le cadre de la mission — soit " + " et ".join(precisions) + " — "
            "nous n'avons connaissance d'aucun contrôle fiscal, "
            "redressement, litige ou passif fiscal, avéré ou éventuel, qui "
            "ne vous aurait pas été signalé.",
            style="List Bullet",
        )
    else:
        doc.add_paragraph(
            "Nous n'avons connaissance d'aucun contrôle fiscal, "
            "redressement, litige ou passif fiscal, avéré ou éventuel, qui "
            "ne vous aurait pas été signalé.",
            style="List Bullet",
        )
    doc.add_paragraph(
        "Tous les passifs fiscaux de l'entité, réels ou potentiels, ont été "
        "comptabilisés ou portés à votre connaissance ; aucun engagement "
        "susceptible d'avoir une incidence fiscale significative n'a été "
        "omis.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Les réponses apportées à vos demandes de renseignements et de "
        "documents sont complètes et sincères ; aucune information "
        "pertinente pour votre mission ne vous a été volontairement "
        "dissimulée.",
        style="List Bullet",
    )

    doc.add_paragraph(
        "Nous vous confirmons ces affirmations en toute connaissance de "
        "leur importance pour les conclusions de votre mission de revue "
        "fiscale."
    )
    doc.add_paragraph(
        "Nous vous prions d'agréer, Madame, Monsieur, l'expression de nos "
        "salutations distinguées."
    )

    # Bloc date/lieu et signature du représentant légal.
    doc.add_paragraph("")
    doc.add_paragraph(
        f"Fait à {_champ(client.get('commune'))}, "
        f"le {date.today().strftime('%d/%m/%Y')}"
    )
    doc.add_paragraph("Le représentant légal")
    doc.add_paragraph(f"Nom et qualité : {A_COMPLETER}")
    doc.add_paragraph("Signature :")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generer_lettre_affirmation_complete(
    session: Session, tenant_id: int, mission_id: int
) -> tuple[bytes, str, dict[str, Any]]:
    """Contenu .docx + nom de fichier + stats — point d'entrée de la route.

    Les stats alimentent le journal d'audit (traçabilité de ce qui a été
    pré-rempli dans la lettre soumise à signature).
    """
    donnees = collecter_donnees_lettre_affirmation(session, tenant_id, mission_id)
    contenu = rendre_lettre_affirmation_docx(donnees)
    nom = nom_fichier_lettre_affirmation(
        donnees["contribuable"].get("denomination"),
        donnees["mission"].get("exercice"),
    )
    stats: dict[str, Any] = {
        "nb_risques_ouverts": int(donnees["nb_risques_ouverts"]),
        "nb_anomalies_derniere_execution": int(donnees["nb_anomalies"]),
    }
    return contenu, nom, stats


def generer_lettre_affirmation(
    session: Session, tenant_id: int, mission_id: int
) -> tuple[bytes, str]:
    """Contenu .docx + nom de fichier — point d'entrée simple.

    Ne lève que si la mission est hors tenant : la lettre est toujours
    produite (pièce du dossier de travail jamais en échec).
    """
    contenu, nom, _stats = generer_lettre_affirmation_complete(
        session, tenant_id, mission_id
    )
    return contenu, nom
