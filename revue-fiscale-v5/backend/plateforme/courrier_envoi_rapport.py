"""Courrier d'envoi du rapport — lettre d'accompagnement client (.docx).

Pratique de cabinet (expertise comptable / conseil fiscal ivoirien) : le
rapport de restitution n'est jamais remis « nu ». Il est accompagné d'une
lettre à en-tête du cabinet qui rappelle la mission et l'exercice, liste
les livrables remis, synthétise les principaux constats chiffrés et
invite le client à une réunion de restitution — signée par l'associé.

Assemblage DÉTERMINISTE (aucun appel LLM, lecture seule sous RLS via
``contexte_tenant``). La synthèse chiffrée provient de la DERNIÈRE
exécution de la mission (mêmes tables execution/conclusion que
``comparatif_executions``) ; s'il n'y a AUCUNE exécution, la lettre est
produite quand même avec la mention « constats en cours d'instruction »
— jamais d'erreur pour ce cas (la pièce du dossier de travail ne doit
jamais être omise). Aucun taux ni seuil fiscal ici — document d'envoi.

Le courrier rappelle en outre les actions du plan d'actions marquées
« retenue » par le fiscaliste (décisions humaines persistées dans
``suivi_plan_actions``, clé ``risque:{id}``) : section « Actions
convenues à mettre en œuvre » — omise s'il n'y a aucune action retenue.
Formulation consultative : le client reste seul décideur.
"""
from __future__ import annotations

import io
import re
import unicodedata
from datetime import date
from decimal import Decimal
from typing import Any, Final

from docx import Document
from sqlalchemy import text
from sqlalchemy.orm import Session

from backend.plateforme.contexte import contexte_tenant
from backend.plateforme.demande_renseignements import A_COMPLETER, _champ
from backend.plateforme.plan_actions import (
    DECISION_RETENUE,
    PREFIXE_CLE_RISQUE,
    _exposition,
)

# Livrables usuellement joints au courrier — liste indicative que
# l'associé raye ou complète à la main avant envoi (pratique cabinet).
LIVRABLES_REMIS: Final[tuple[str, ...]] = (
    "Rapport de revue fiscale (version Word et version PDF)",
    "Note de synthèse de mission à l'attention de la Direction",
    "Plan d'actions correctives et recommandations",
    "Annexes chiffrées (détail des constats par impôt)",
)


class ErreurCourrierEnvoi(Exception):
    """Echec de génération du courrier d'envoi du rapport."""


class ErreurCourrierEnvoiIntrouvable(ErreurCourrierEnvoi):
    """Mission hors périmètre du tenant — 404 côté route."""


def _fmt_montant(montant: object) -> str:
    """« 1 234 567 » — séparateur d'espace, sans décimales (usage FCFA)."""
    try:
        return f"{Decimal(str(montant)):,.0f}".replace(",", " ")
    except Exception:  # noqa: BLE001 — valeur non numérique tolérée
        return str(montant)


def nom_fichier_courrier_envoi(
    denomination: object | None, exercice: object | None
) -> str:
    """courrier_envoi_rapport_{NOM}_{exercice}.docx — ASCII sûr (HTTP)."""
    brut = str(denomination or "client")
    sans_accents = (
        unicodedata.normalize("NFKD", brut).encode("ascii", "ignore").decode("ascii")
    )
    nom = re.sub(r"[^A-Za-z0-9]+", "_", sans_accents).strip("_").upper() or "CLIENT"
    exo = str(exercice or A_COMPLETER).strip() or "exercice"
    exo = re.sub(r"[^A-Za-z0-9]+", "_", exo) or "exercice"
    return f"courrier_envoi_rapport_{nom}_{exo}.docx"


def _synthese_derniere_execution(
    session: Session, mission_id: int
) -> dict[str, Any] | None:
    """Compteurs de la DERNIÈRE exécution — None si la mission n'en a pas.

    Mêmes tables que ``comparatif_executions`` : conclusion jointe à
    l'exécution la plus récente. Total des montants d'anomalies en
    ``Decimal`` (jamais de float sur un montant).
    """
    execution = session.execute(
        text(
            "SELECT id, lancee_le FROM execution "
            "WHERE mission_id = :m ORDER BY id DESC LIMIT 1"
        ),
        {"m": mission_id},
    ).mappings().one_or_none()
    if execution is None:
        return None
    rows = session.execute(
        text(
            "SELECT c.statut, c.montant FROM conclusion c "
            "WHERE c.execution_id = :e"
        ),
        {"e": int(execution["id"])},
    ).mappings().all()
    conformes = anomalies = non_verifiables = 0
    total_anomalies = Decimal("0")
    for r in rows:
        statut = str(r["statut"] or "anomalie")
        if statut == "conforme":
            conformes += 1
        elif statut == "anomalie":
            anomalies += 1
            if r["montant"] is not None:
                try:
                    total_anomalies += Decimal(str(r["montant"]))
                except Exception:  # noqa: BLE001 — montant non numérique
                    pass
        elif statut == "non_verifiable":
            non_verifiables += 1
    lancee_le = execution.get("lancee_le")
    return {
        "execution_id": int(execution["id"]),
        "lancee_le": (
            lancee_le.isoformat() if hasattr(lancee_le, "isoformat") else lancee_le
        ),
        "conformes": conformes,
        "anomalies": anomalies,
        "non_verifiables": non_verifiables,
        "total_constats": len(rows),
        "total_montant_anomalies": total_anomalies,
    }


def composer_lignes_actions_convenues(
    actions: list[dict[str, Any]],
) -> list[str]:
    """PUR — une ligne de puce par action retenue du plan d'actions.

    ``actions`` : items {libelle_risque, impot, exposition, decision_note}
    (sortie de :func:`_actions_retenues_mission`). Libellé du risque,
    impôt et exposition estimée en FCFA quand ils sont connus, note de
    décision du fiscaliste si présente. Liste vide → liste vide (le
    courrier n'affiche alors pas la section).
    """
    lignes: list[str] = []
    for a in actions:
        libelle = str(a.get("libelle_risque") or "").strip() or (
            "Action retenue au plan d'actions"
        )
        morceaux = [libelle]
        impot = str(a.get("impot") or "").strip()
        if impot:
            morceaux.append(f"impôt : {impot.upper()}")
        exposition = a.get("exposition")
        if exposition is not None and str(exposition).strip() != "":
            morceaux.append(
                f"exposition estimée : {_fmt_montant(exposition)} FCFA"
            )
        note = str(a.get("decision_note") or "").strip()
        if note:
            morceaux.append(f"note : {note}")
        lignes.append(" — ".join(morceaux))
    return lignes


def _actions_retenues_mission(
    session: Session, mission_id: int
) -> list[dict[str, Any]]:
    """Actions « retenues » (non faites) de la mission — décisions humaines.

    Table ``suivi_plan_actions`` (UPSERT par ``cle_action`` : une décision
    ultérieure « faite » / « écartée » remplace la ligne, seul l'état
    courant « retenue » est donc listé). Le risque d'origine est rejoint
    via ``cle_action`` (``risque:{id}``) pour le libellé, l'impôt et
    l'exposition. Contexte tenant déjà posé par l'appelant.
    """
    rows = session.execute(
        text(
            "SELECT s.cle_action, s.note, r.libelle, r.impot, "
            "r.montant_estime, r.penalites_estimees "
            "FROM suivi_plan_actions s "
            "JOIN mission m ON m.id = s.mission_id "
            "LEFT JOIN risque r "
            "ON s.cle_action = :prefixe || r.id::text "
            "AND r.contribuable_id = m.contribuable_id "
            "WHERE s.mission_id = :m AND s.decision = :d "
            "ORDER BY s.id ASC"
        ),
        {"m": mission_id, "d": DECISION_RETENUE, "prefixe": PREFIXE_CLE_RISQUE},
    ).mappings().all()
    actions: list[dict[str, Any]] = []
    for r in rows:
        exposition = _exposition(dict(r))
        actions.append(
            {
                "cle_action": str(r["cle_action"]),
                "libelle_risque": str(r["libelle"] or ""),
                "impot": str(r["impot"] or "").upper(),
                "exposition": (
                    str(exposition) if exposition is not None else None
                ),
                "decision_note": (r["note"] or None) or None,
            }
        )
    return actions


def collecter_donnees_courrier_envoi(
    session: Session, tenant_id: int, mission_id: int
) -> dict[str, Any]:
    """Lecture seule (RLS via contexte_tenant) : mission + client + synthèse.

    Mission hors tenant → :class:`ErreurCourrierEnvoiIntrouvable` (404).
    L'identité du cabinet (table tenant, sans RLS) est lue par
    id = tenant_id — même garde que /api/v1/auth/connexion.
    """
    with contexte_tenant(session, tenant_id):
        row = session.execute(
            text(
                "SELECT m.id, m.exercice, m.statut, "
                "c.denomination AS contribuable_denomination, c.ncc, "
                "c.siege_social, c.commune "
                "FROM mission m JOIN contribuable c ON c.id = m.contribuable_id "
                "WHERE m.id = :m"
            ),
            {"m": mission_id},
        ).mappings().one_or_none()
        if row is None:
            raise ErreurCourrierEnvoiIntrouvable(
                f"mission {mission_id} introuvable"
            )
        synthese = _synthese_derniere_execution(session, mission_id)
        actions_retenues = _actions_retenues_mission(session, mission_id)

    cabinet = session.execute(
        text(
            "SELECT denomination, ncc, rccm, forme_juridique, siege_social, "
            "commune, centre_impots "
            "FROM tenant WHERE id = :t"
        ),
        {"t": tenant_id},
    ).mappings().one_or_none()

    return {
        "mission": {
            "id": int(row["id"]),
            "exercice": row["exercice"],
            "statut": row["statut"],
        },
        "contribuable": {
            "denomination": row["contribuable_denomination"],
            "ncc": row["ncc"],
            "siege_social": row["siege_social"],
            "commune": row["commune"],
        },
        "cabinet": dict(cabinet) if cabinet is not None else {},
        "synthese": synthese,
        "actions_retenues": actions_retenues,
    }


def rendre_courrier_envoi_docx(donnees: dict[str, Any]) -> bytes:
    """Assemble le .docx — en-tête cabinet, synthèse chiffrée, signature."""
    mission: dict[str, Any] = donnees.get("mission") or {}
    client: dict[str, Any] = donnees.get("contribuable") or {}
    cabinet: dict[str, Any] = donnees.get("cabinet") or {}
    synthese: dict[str, Any] | None = donnees.get("synthese")

    exercice = mission.get("exercice")
    doc = Document()

    # En-tête cabinet (émetteur) — même style que les autres courriers.
    doc.add_paragraph(_champ(cabinet.get("denomination")).upper())
    doc.add_paragraph(
        f"Forme juridique : {_champ(cabinet.get('forme_juridique'))} — "
        f"RCCM : {_champ(cabinet.get('rccm'))} — NCC : {_champ(cabinet.get('ncc'))}"
    )
    doc.add_paragraph(
        f"Siège : {_champ(cabinet.get('siege_social'))} — "
        f"{_champ(cabinet.get('commune'))}"
    )
    doc.add_paragraph(
        f"Centre des impôts de rattachement : {_champ(cabinet.get('centre_impots'))}"
    )

    # Destinataire (contribuable, NCC seulement s'il est connu).
    doc.add_paragraph("")
    doc.add_paragraph(
        f"À l'attention de la Direction de {_champ(client.get('denomination'))}"
    )
    ncc = str(client.get("ncc") or "").strip()
    if ncc:
        doc.add_paragraph(f"NCC : {ncc}")
    doc.add_paragraph(
        f"Siège : {_champ(client.get('siege_social'))} — "
        f"{_champ(client.get('commune'))}"
    )
    doc.add_paragraph(
        f"{_champ(cabinet.get('commune'))}, le {date.today().strftime('%d/%m/%Y')}"
    )

    # Objet
    doc.add_heading("Courrier d'envoi du rapport de revue fiscale", level=1)
    doc.add_paragraph(
        f"Objet : remise du rapport de revue fiscale — exercice "
        f"{_champ(exercice)}."
    )
    doc.add_paragraph("Madame, Monsieur,")

    # Rappel de la mission.
    doc.add_paragraph(
        "Au terme de nos travaux, nous avons le plaisir de vous remettre, "
        "sous ce pli, le rapport de notre mission de revue fiscale "
        f"(mission n° {mission.get('id')}) portant sur l'exercice "
        f"{_champ(exercice)}. À la date du présent courrier, la mission est "
        f"au statut « {_champ(mission.get('statut'))} » dans nos dossiers."
    )

    # Livrables remis.
    doc.add_heading("Documents remis", level=2)
    doc.add_paragraph(
        "Vous trouverez joints au présent courrier les livrables suivants "
        "(liste à ajuster selon les documents effectivement remis) :"
    )
    for livrable in LIVRABLES_REMIS:
        doc.add_paragraph(livrable, style="List Bullet")

    # Synthèse chiffrée des constats — dernière exécution, ou mention
    # « en cours d'instruction » si la mission n'a pas encore été exécutée.
    doc.add_heading("Principaux constats", level=2)
    if synthese is None:
        doc.add_paragraph(
            "Les constats de la mission sont en cours d'instruction à la "
            "date d'édition du présent courrier : la synthèse chiffrée vous "
            "sera communiquée avec la version définitive du rapport."
        )
    else:
        doc.add_paragraph(
            f"Notre revue a porté sur {synthese['total_constats']} point(s) "
            "de contrôle. Il en ressort :"
        )
        doc.add_paragraph(
            f"{synthese['conformes']} constat(s) conforme(s) ;",
            style="List Bullet",
        )
        doc.add_paragraph(
            f"{synthese['anomalies']} anomalie(s), pour un enjeu total "
            f"estimé à {_fmt_montant(synthese['total_montant_anomalies'])} "
            "FCFA ;",
            style="List Bullet",
        )
        doc.add_paragraph(
            f"{synthese['non_verifiables']} constat(s) non vérifiable(s) en "
            "l'état des pièces communiquées.",
            style="List Bullet",
        )
        doc.add_paragraph(
            "Le détail de chaque constat, ses références légales et nos "
            "recommandations figurent dans le rapport joint."
        )

    # Actions retenues du plan d'actions — décisions humaines persistées
    # (suivi_plan_actions). Aucune action retenue → pas de section.
    actions_retenues: list[dict[str, Any]] = (
        donnees.get("actions_retenues") or []
    )
    if actions_retenues:
        doc.add_heading("Actions convenues à mettre en œuvre", level=2)
        doc.add_paragraph(
            "À l'issue de nos échanges sur le plan d'actions, les actions "
            "suivantes ont été retenues d'un commun accord et restent à "
            "mettre en œuvre :"
        )
        for ligne in composer_lignes_actions_convenues(actions_retenues):
            doc.add_paragraph(ligne, style="List Bullet")
        doc.add_paragraph(
            "Ces actions constituent des recommandations de notre cabinet : "
            "la décision de leur mise en œuvre, ainsi que son calendrier, "
            "relèvent de votre seule appréciation. Nous restons à vos côtés "
            "pour vous accompagner dans leur réalisation si vous le "
            "souhaitez."
        )

    # Invitation à la réunion de restitution — pratique cabinet.
    doc.add_heading("Réunion de restitution", level=2)
    doc.add_paragraph(
        "Nous vous proposons d'organiser une réunion de restitution afin de "
        "vous présenter ces conclusions, de recueillir vos observations et "
        "de convenir ensemble des suites à donner. Nous vous remercions de "
        "bien vouloir nous indiquer vos disponibilités. "
        f"Date proposée : {A_COMPLETER}."
    )

    # Formule de politesse + signature de l'associé.
    doc.add_paragraph(
        "Nous restons à votre entière disposition pour tout complément "
        "d'information et vous prions d'agréer, Madame, Monsieur, "
        "l'expression de nos salutations distinguées."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"Pour le cabinet : {_champ(cabinet.get('denomination'))}")
    doc.add_paragraph("L'associé signataire")
    doc.add_paragraph("Nom et qualité : [à compléter]")
    doc.add_paragraph("Signature :")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generer_courrier_envoi_complet(
    session: Session, tenant_id: int, mission_id: int
) -> tuple[bytes, str, dict[str, Any]]:
    """Contenu .docx + nom de fichier + stats — point d'entrée de la route.

    Les stats alimentent le journal d'audit (traçabilité de ce qui a été
    annoncé au client dans le courrier).
    """
    donnees = collecter_donnees_courrier_envoi(session, tenant_id, mission_id)
    contenu = rendre_courrier_envoi_docx(donnees)
    nom = nom_fichier_courrier_envoi(
        donnees["contribuable"].get("denomination"),
        donnees["mission"].get("exercice"),
    )
    synthese = donnees.get("synthese")
    stats: dict[str, Any] = {
        "avec_execution": synthese is not None,
        "nb_conformes": synthese["conformes"] if synthese else 0,
        "nb_anomalies": synthese["anomalies"] if synthese else 0,
        "nb_non_verifiables": synthese["non_verifiables"] if synthese else 0,
        "total_montant_anomalies": (
            str(synthese["total_montant_anomalies"]) if synthese else "0"
        ),
        "nb_actions_retenues": len(donnees.get("actions_retenues") or []),
    }
    return contenu, nom, stats


def generer_courrier_envoi(
    session: Session, tenant_id: int, mission_id: int
) -> bytes:
    """Bytes du .docx — point d'entrée simple (pièce du dossier de travail).

    Ne lève que si la mission est hors tenant : sans exécution, la lettre
    est produite avec la mention « constats en cours d'instruction ».
    """
    contenu, _nom, _stats = generer_courrier_envoi_complet(
        session, tenant_id, mission_id
    )
    return contenu
