"""Courrier de relance — demande de renseignements et de documents (.docx).

Quand le client tarde à répondre à la « Demande de renseignements et de
documents », le cabinet lui adresse un courrier de relance listant les
SEULS items encore en attente (suivi de circularisation :
``suivi_renseignements.lister_items``). Assemblage DÉTERMINISTE (aucun
appel LLM) — même style d'en-tête cabinet/client que
``demande_renseignements``. Aucun taux ni seuil fiscal ici.
"""
from __future__ import annotations

import io
import re
import unicodedata
from datetime import date, timedelta
from typing import Any, Final

from docx import Document
from sqlalchemy import text
from sqlalchemy.orm import Session

from backend.plateforme.contexte import contexte_tenant
from backend.plateforme.demande_renseignements import (
    A_COMPLETER,
    DELAI_REPONSE_JOURS,
    _champ,
)
from backend.plateforme.suivi_renseignements import (
    STATUT_DEFAUT,
    lister_items,
)

DELAI_RELANCE_JOURS: Final = 8

#: Délai de réponse « sous quinzaine » du courrier texte (jours calendaires).
DELAI_QUINZAINE_JOURS: Final = 15

MENTION_COURRIER_TXT: Final = (
    "Courrier généré automatiquement à partir de la demande de "
    "renseignements — à relire et adapter par le fiscaliste avant envoi."
)


class ErreurCourrierRelance(Exception):
    """Echec de génération du courrier de relance."""


class ErreurCourrierIntrouvable(ErreurCourrierRelance):
    """Mission hors périmètre du tenant — 404 côté route."""


class ErreurAucunItemEnAttente(ErreurCourrierRelance):
    """Aucun item en attente : la relance n'a pas d'objet — 409 côté route."""


def _date_fr(iso: object | None) -> str | None:
    """« JJ/MM/AAAA » depuis une date ISO — None si absente/invalide."""
    brut = str(iso or "").strip()
    if not brut:
        return None
    try:
        return date.fromisoformat(brut[:10]).strftime("%d/%m/%Y")
    except ValueError:
        return None


def nom_fichier_relance(
    denomination: object | None, exercice: object | None
) -> str:
    """relance_{NOM}_{exercice}.docx — nom ASCII sûr (HTTP)."""
    brut = str(denomination or "client")
    sans_accents = (
        unicodedata.normalize("NFKD", brut).encode("ascii", "ignore").decode("ascii")
    )
    nom = re.sub(r"[^A-Za-z0-9]+", "_", sans_accents).strip("_").upper() or "CLIENT"
    exo = str(exercice or A_COMPLETER).strip() or "exercice"
    exo = re.sub(r"[^A-Za-z0-9]+", "_", exo) or "exercice"
    return f"relance_{nom}_{exo}.docx"


def collecter_donnees_relance(
    session: Session, tenant_id: int, mission_id: int
) -> dict[str, Any]:
    """Lecture seule (RLS via contexte_tenant) : mission + client + items.

    Seuls les items du suivi en statut ``en_attente`` sont retenus. Mission
    hors tenant → :class:`ErreurCourrierIntrouvable` ; aucun item en
    attente → :class:`ErreurAucunItemEnAttente`. L'identité du cabinet
    (table tenant, sans RLS) est lue par id = tenant_id, même garde que
    /api/v1/auth/connexion.
    """
    from backend.plateforme.suivi_renseignements import ErreurSuiviIntrouvable

    try:
        items = lister_items(session, tenant_id, mission_id)
    except ErreurSuiviIntrouvable as e:
        raise ErreurCourrierIntrouvable(str(e)) from e
    en_attente = [
        i for i in items if str(i.get("statut") or "") == STATUT_DEFAUT
    ]
    if not en_attente:
        raise ErreurAucunItemEnAttente(
            f"aucun item en attente pour la mission {mission_id} — "
            "le courrier de relance est sans objet"
        )

    with contexte_tenant(session, tenant_id):
        row = session.execute(
            text(
                "SELECT m.id, m.exercice, "
                "c.denomination AS contribuable_denomination, c.ncc, "
                "c.siege_social, c.commune "
                "FROM mission m JOIN contribuable c ON c.id = m.contribuable_id "
                "WHERE m.id = :m"
            ),
            {"m": mission_id},
        ).mappings().one_or_none()
        if row is None:  # défense en profondeur — lister_items a déjà vérifié
            raise ErreurCourrierIntrouvable(f"mission {mission_id} introuvable")

    cabinet = session.execute(
        text(
            "SELECT denomination, ncc, rccm, forme_juridique, siege_social, "
            "commune, centre_impots "
            "FROM tenant WHERE id = :t"
        ),
        {"t": tenant_id},
    ).mappings().one_or_none()

    return {
        "mission": {"id": int(row["id"]), "exercice": row["exercice"]},
        "contribuable": {
            "denomination": row["contribuable_denomination"],
            "ncc": row["ncc"],
            "siege_social": row["siege_social"],
            "commune": row["commune"],
        },
        "cabinet": dict(cabinet) if cabinet is not None else {},
        "items_en_attente": en_attente,
        "nb_items_total": len(items),
    }


def rendre_relance_docx(donnees: dict[str, Any]) -> bytes:
    """Assemble le .docx — items en attente numérotés (1., 2., …)."""
    mission: dict[str, Any] = donnees.get("mission") or {}
    client: dict[str, Any] = donnees.get("contribuable") or {}
    cabinet: dict[str, Any] = donnees.get("cabinet") or {}
    items: list[dict[str, Any]] = donnees.get("items_en_attente") or []

    exercice = mission.get("exercice")
    doc = Document()

    # En-tête cabinet (émetteur) — même style que la demande de renseignements.
    doc.add_paragraph(_champ(cabinet.get("denomination")).upper())
    doc.add_paragraph(
        f"Forme juridique : {_champ(cabinet.get('forme_juridique'))} — "
        f"RCCM : {_champ(cabinet.get('rccm'))} — NCC : {_champ(cabinet.get('ncc'))}"
    )
    doc.add_paragraph(
        f"Siège : {_champ(cabinet.get('siege_social'))} — "
        f"{_champ(cabinet.get('commune'))}"
    )
    doc.add_paragraph(
        f"Centre des impôts de rattachement : {_champ(cabinet.get('centre_impots'))}"
    )

    # Destinataire
    doc.add_paragraph("")
    doc.add_paragraph(
        f"À l'attention de la Direction de {_champ(client.get('denomination'))}"
    )
    doc.add_paragraph(
        f"Siège : {_champ(client.get('siege_social'))} — "
        f"{_champ(client.get('commune'))}"
    )
    doc.add_paragraph(
        f"{_champ(cabinet.get('commune'))}, le {date.today().strftime('%d/%m/%Y')}"
    )

    # Objet
    doc.add_heading("Relance — demande de renseignements et de documents", level=1)
    doc.add_paragraph(
        f"Objet : relance — mission de revue fiscale — exercice "
        f"{_champ(exercice)} — demande de renseignements et de documents."
    )
    doc.add_paragraph("Madame, Monsieur,")

    # Rappel courtois : demande initiale + délai indicatif dépassé.
    doc.add_paragraph(
        "Dans le cadre de notre mission de revue fiscale de l'exercice "
        f"{_champ(exercice)}, nous vous avons adressé une demande de "
        "renseignements et de documents. Sauf erreur ou omission de notre "
        f"part, le délai indicatif de {DELAI_REPONSE_JOURS} jours mentionné "
        "dans ce courrier est désormais dépassé et certains éléments "
        "demeurent en attente."
    )

    # Liste numérotée des SEULS items en attente.
    doc.add_heading("Éléments toujours en attente", level=2)
    doc.add_paragraph(
        "Nous nous permettons de rappeler ci-après les éléments toujours en "
        "attente à la date d'édition du présent courrier. Chaque élément est "
        "numéroté : merci de rappeler ce numéro dans votre réponse."
    )
    for numero, it in enumerate(items, start=1):
        libelle = str(it.get("libelle") or "").strip() or A_COMPLETER
        relance = _date_fr(it.get("date_relance"))
        suffixe = f" (relance prévue le {relance})" if relance else ""
        doc.add_paragraph(f"{numero}. {libelle}{suffixe}")

    # Clôture : nouveau délai, conséquences, coordonnées du cabinet.
    doc.add_heading("Suite à donner", level=2)
    doc.add_paragraph(
        "Nous vous remercions de nous faire parvenir vos réponses et les "
        "pièces demandées sous un nouveau délai indicatif de "
        f"{DELAI_RELANCE_JOURS} jours à compter de la réception de la "
        "présente relance."
    )
    doc.add_paragraph(
        "À défaut de réponse dans ce délai, les constats concernés seront "
        "maintenus « non vérifiables » dans notre rapport et des réserves "
        "pourront être formulées sur les points correspondants."
    )
    doc.add_paragraph(
        f"Contact du cabinet : {_champ(cabinet.get('denomination'))} — "
        f"{_champ(cabinet.get('siege_social'))} — "
        f"{_champ(cabinet.get('commune'))}. "
        f"Interlocuteur : {A_COMPLETER}."
    )
    doc.add_paragraph(
        "Nous restons à votre disposition pour toute précision et vous "
        "prions d'agréer, Madame, Monsieur, l'expression de nos "
        "salutations distinguées."
    )
    doc.add_paragraph(f"Pour le cabinet : {_champ(cabinet.get('denomination'))}")
    doc.add_paragraph("Nom et qualité : [à compléter]")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generer_courrier_relance(
    session: Session, tenant_id: int, mission_id: int
) -> tuple[bytes, str, dict[str, int]]:
    """Contenu .docx + nom de fichier + compteurs — point d'entrée route."""
    donnees = collecter_donnees_relance(session, tenant_id, mission_id)
    contenu = rendre_relance_docx(donnees)
    nom = nom_fichier_relance(
        donnees["contribuable"].get("denomination"),
        donnees["mission"].get("exercice"),
    )
    stats = {
        "nb_items_en_attente": len(donnees["items_en_attente"]),
        "nb_items_total": int(donnees["nb_items_total"]),
    }
    return contenu, nom, stats


# ── Courrier de relance texte brut (.txt) ────────────────────────────


def construire_courrier(contexte: dict[str, Any]) -> str:
    """PUR — courrier de relance en texte brut, entièrement déterministe.

    ``contexte`` : {cabinet, contribuable, exercice, aujourd_hui (date),
    items ([{libelle, date_relance?, nb_relances?, derniere_relance_le?}])}.
    Items numérotés 1., 2., … avec date de demande/relance « JJ/MM/AAAA »
    si disponible. Historique de circularisation (migration 042) : rang de
    relance global = max(nb_relances des items ouverts) + 1 ; à partir du
    2e rang l'objet devient « 2e relance — … » et une phrase rappelle la
    dernière relance ; chaque item déjà relancé porte « (déjà relancé
    N fois) ». Clôture courtoise « sous quinzaine » (date limite =
    ``aujourd_hui`` + 15 jours calendaires). Sans item : courrier
    signalant qu'aucune relance n'est nécessaire. La date du jour vient
    du paramètre ``aujourd_hui`` (aucun ``date.today()`` ici).
    """
    cabinet = _champ(contexte.get("cabinet"))
    contribuable = _champ(contexte.get("contribuable"))
    exercice = _champ(contexte.get("exercice"))
    jour: date = contexte["aujourd_hui"]
    items: list[dict[str, Any]] = list(contexte.get("items") or [])

    # Rang de relance global : max(nb_relances des items ouverts) + 1.
    rang = max(
        (int(it.get("nb_relances") or 0) for it in items), default=0
    ) + 1
    # Date de la dernière relance passée = max des derniere_relance_le
    # non nulles (comparaison ISO = comparaison chronologique).
    derniere_relance = _date_fr(
        max(
            (
                str(it.get("derniere_relance_le") or "").strip()
                for it in items
                if str(it.get("derniere_relance_le") or "").strip()
            ),
            default="",
        )
        or None
    )
    objet = "Relance" if rang < 2 else f"{rang}e relance"

    lignes: list[str] = [
        cabinet.upper(),
        f"Le {jour.strftime('%d/%m/%Y')}",
        "",
        f"À l'attention de la Direction de {contribuable}",
        "",
        f"Objet : {objet} — pièces et renseignements en attente "
        f"(mission {exercice})",
        "",
        "Madame, Monsieur,",
        "",
    ]
    if items:
        if rang >= 2 and derniere_relance:
            lignes += [
                f"Nous vous avons déjà relancés le {derniere_relance}. "
                "Sauf erreur ou omission de notre part, les éléments "
                "rappelés ci-dessous demeurent toutefois en attente.",
                "",
            ]
        lignes += [
            "Dans le cadre de notre mission de revue fiscale de "
            f"l'exercice {exercice}, les éléments suivants, demandés dans "
            "notre demande de renseignements et de documents, demeurent en "
            "attente à ce jour :",
            "",
        ]
        for numero, it in enumerate(items, start=1):
            libelle = str(it.get("libelle") or "").strip() or A_COMPLETER
            date_demande = _date_fr(it.get("date_relance"))
            suffixe = f" (demande du {date_demande})" if date_demande else ""
            nb = int(it.get("nb_relances") or 0)
            if nb >= 1:
                suffixe += f" (déjà relancé {nb} fois)"
            lignes.append(f"{numero}. {libelle}{suffixe}")
        date_limite = (jour + timedelta(days=DELAI_QUINZAINE_JOURS)).strftime(
            "%d/%m/%Y"
        )
        lignes += [
            "",
            "Nous vous remercions de bien vouloir nous faire parvenir ces "
            "éléments dans les meilleurs délais.",
            "",
            "Nous vous saurions gré de nous répondre sous quinzaine, soit "
            f"au plus tard le {date_limite} (15 jours calendaires à "
            "compter de la date du présent courrier).",
        ]
    else:
        lignes.append(
            "Aucun élément n'est en attente à ce jour : aucune relance "
            "n'est nécessaire."
        )
    lignes += [
        "",
        "Nous vous prions d'agréer, Madame, Monsieur, l'expression de nos "
        "salutations distinguées.",
        "",
        f"Pour le cabinet : {cabinet}",
        "",
        MENTION_COURRIER_TXT,
    ]
    return "\n".join(lignes) + "\n"


def courrier_mission(
    session: Session,
    tenant_id: int,
    mission_id: int,
    *,
    aujourd_hui: date | None = None,
) -> dict[str, Any]:
    """Courrier de relance texte de la mission (lecture seule, RLS).

    Items OUVERTS = statut ``en_attente`` du suivi de circularisation
    (les ``recu`` / ``sans_objet`` sont soldés). Mission hors tenant →
    :class:`ErreurCourrierIntrouvable` (404 côté route). Sans item
    ouvert : ``nb_items_ouverts = 0`` et courrier signalant qu'aucune
    relance n'est nécessaire (pas d'erreur).
    """
    from backend.plateforme.suivi_renseignements import (
        ErreurSuiviIntrouvable,
    )

    jour = aujourd_hui or date.today()
    # lister_items ouvre son propre contexte_tenant : appel HORS de tout
    # autre with contexte_tenant.
    try:
        items = lister_items(session, tenant_id, mission_id)
    except ErreurSuiviIntrouvable as e:
        raise ErreurCourrierIntrouvable(str(e)) from e
    ouverts = [
        i for i in items if str(i.get("statut") or "") == STATUT_DEFAUT
    ]

    with contexte_tenant(session, tenant_id):
        row = session.execute(
            text(
                "SELECT m.exercice, c.denomination "
                "FROM mission m JOIN contribuable c "
                "ON c.id = m.contribuable_id WHERE m.id = :m"
            ),
            {"m": mission_id},
        ).mappings().one_or_none()
        if row is None:  # défense en profondeur — lister_items a vérifié
            raise ErreurCourrierIntrouvable(
                f"mission {mission_id} introuvable"
            )

    # Identité du cabinet (table tenant, sans RLS) — même garde que
    # /api/v1/auth/connexion.
    cabinet = session.execute(
        text("SELECT denomination FROM tenant WHERE id = :t"),
        {"t": tenant_id},
    ).scalar_one_or_none()

    courrier = construire_courrier(
        {
            "cabinet": cabinet,
            "contribuable": row["denomination"],
            "exercice": row["exercice"],
            "aujourd_hui": jour,
            "items": ouverts,
        }
    )
    return {
        "mission_id": mission_id,
        "contribuable": row["denomination"],
        "exercice": row["exercice"],
        "nb_items_ouverts": len(ouverts),
        "courrier": courrier,
        "note": MENTION_COURRIER_TXT,
    }
