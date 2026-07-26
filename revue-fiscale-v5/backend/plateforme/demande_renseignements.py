"""Demande de renseignements et de documents — livrable Word (.docx).

Après les travaux de revue, l'expert adresse au client une demande de
renseignements et de documents (circularisation) listant les questions et
pièces à fournir. Assemblage DÉTERMINISTE (aucun appel LLM) depuis :

- la dernière version « disponible » du commentaire IA de revue analytique
  (``commentaire_revue_analytique``) : questions à poser au client ;
- les conclusions ``non_verifiable`` de la dernière exécution de la
  mission (mêmes tables/colonnes que ``note_synthese.construire_contexte``
  via ``restitution.service._charger_conclusions``) : pièces / réponses
  manquantes, chacune citant sa règle (``regle_id``).

Les demandes sont numérotées en continu (1., 2., …) pour faciliter la
réponse du client. Aucun taux ni seuil fiscal ici — document de liaison.
"""
from __future__ import annotations

import io
import re
import unicodedata
from datetime import date
from typing import Any, Final

from docx import Document
from sqlalchemy import text
from sqlalchemy.orm import Session

from backend.plateforme.contexte import contexte_tenant

A_COMPLETER: Final = "[à compléter]"
STATUT_NON_VERIFIABLE: Final = "non_verifiable"
DELAI_REPONSE_JOURS: Final = 15


class ErreurDemandeRenseignements(Exception):
    """Echec de génération de la demande (mission introuvable…)."""


def _champ(valeur: object | None) -> str:
    """Valeur affichable — jamais d'invention : manquant = « [à compléter] »."""
    if valeur is None:
        return A_COMPLETER
    texte_ = str(valeur).strip()
    return texte_ or A_COMPLETER


def nom_fichier_demande(
    denomination: object | None, exercice: object | None
) -> str:
    """demande_renseignements_{NOM}_{exercice}.docx — nom ASCII sûr (HTTP)."""
    brut = str(denomination or "client")
    sans_accents = (
        unicodedata.normalize("NFKD", brut).encode("ascii", "ignore").decode("ascii")
    )
    nom = re.sub(r"[^A-Za-z0-9]+", "_", sans_accents).strip("_").upper() or "CLIENT"
    exo = str(exercice or A_COMPLETER).strip() or "exercice"
    exo = re.sub(r"[^A-Za-z0-9]+", "_", exo) or "exercice"
    return f"demande_renseignements_{nom}_{exo}.docx"


def _dernier_commentaire_disponible(
    session: Session, mission_id: int
) -> dict[str, Any] | None:
    """Contenu jsonb de la dernière version « disponible » du commentaire.

    Contexte tenant déjà posé par l'appelant. ``None`` si aucun commentaire
    disponible (section analytique alors omise du document).
    """
    contenu = session.execute(
        text(
            "SELECT contenu FROM commentaire_revue_analytique "
            "WHERE mission_id = :m AND statut = 'disponible' "
            "ORDER BY version DESC, id DESC LIMIT 1"
        ),
        {"m": mission_id},
    ).scalar_one_or_none()
    return dict(contenu) if isinstance(contenu, dict) else None


def _conclusions_non_verifiables(
    session: Session, mission_id: int
) -> list[dict[str, Any]]:
    """Conclusions ``non_verifiable`` de la dernière exécution de la mission.

    Mêmes tables/colonnes que ``note_synthese.construire_contexte``
    (``execution`` la plus récente puis ``conclusion`` JOIN
    ``regle_version``), enrichies du libellé de la règle (``regle``) pour
    l'intitulé de la demande. Contexte tenant déjà posé par l'appelant.
    """
    exec_id = session.execute(
        text(
            "SELECT id FROM execution WHERE mission_id = :m "
            "ORDER BY id DESC LIMIT 1"
        ),
        {"m": mission_id},
    ).scalar_one_or_none()
    if exec_id is None:
        return []
    rows = session.execute(
        text(
            "SELECT rv.regle_id, c.commentaire, c.niveau_risque, "
            "r.libelle AS regle_libelle "
            "FROM conclusion c "
            "JOIN regle_version rv ON rv.id = c.regle_version_id "
            "LEFT JOIN regle r ON r.identifiant = rv.regle_id "
            "WHERE c.execution_id = :e AND c.statut = :st "
            "ORDER BY rv.regle_id"
        ),
        {"e": int(exec_id), "st": STATUT_NON_VERIFIABLE},
    ).mappings().all()
    return [
        {
            "regle_id": str(r["regle_id"]),
            "libelle": r.get("regle_libelle"),
            "commentaire": r.get("commentaire"),
            "niveau_risque": r.get("niveau_risque"),
        }
        for r in rows
    ]


def collecter_donnees_demande(
    session: Session, tenant_id: int, mission_id: int
) -> dict[str, Any]:
    """Lecture seule (RLS via contexte_tenant) : mission + client + sources.

    L'identité du cabinet (table tenant, sans RLS) est lue par
    id = tenant_id, même garde que /api/v1/auth/connexion.
    """
    with contexte_tenant(session, tenant_id):
        row = session.execute(
            text(
                "SELECT m.id, m.exercice, "
                "c.denomination AS contribuable_denomination, c.ncc, "
                "c.siege_social, c.commune "
                "FROM mission m JOIN contribuable c ON c.id = m.contribuable_id "
                "WHERE m.id = :m"
            ),
            {"m": mission_id},
        ).mappings().one_or_none()
        if row is None:
            raise ErreurDemandeRenseignements(
                f"mission {mission_id} introuvable"
            )
        commentaire = _dernier_commentaire_disponible(session, mission_id)
        conclusions = _conclusions_non_verifiables(session, mission_id)

    cabinet = session.execute(
        text(
            "SELECT denomination, ncc, rccm, forme_juridique, siege_social, "
            "commune, centre_impots "
            "FROM tenant WHERE id = :t"
        ),
        {"t": tenant_id},
    ).mappings().one_or_none()

    questions: list[dict[str, Any]] = []
    if commentaire is not None:
        for item in commentaire.get("explications") or []:
            if not isinstance(item, dict):
                continue
            question = str(
                item.get("question_a_poser_au_client") or ""
            ).strip()
            if not question:
                continue
            questions.append(
                {
                    "poste": str(item.get("poste") or "").strip(),
                    "question": question,
                    "gravite": str(item.get("gravite") or "").strip(),
                }
            )

    return {
        "mission": {"id": int(row["id"]), "exercice": row["exercice"]},
        "contribuable": {
            "denomination": row["contribuable_denomination"],
            "ncc": row["ncc"],
            "siege_social": row["siege_social"],
            "commune": row["commune"],
        },
        "cabinet": dict(cabinet) if cabinet is not None else {},
        "questions_analytique": questions,
        "conclusions_non_verifiables": conclusions,
    }


def rendre_demande_docx(donnees: dict[str, Any]) -> bytes:
    """Assemble le .docx — demandes numérotées en continu (1., 2., …)."""
    mission: dict[str, Any] = donnees.get("mission") or {}
    client: dict[str, Any] = donnees.get("contribuable") or {}
    cabinet: dict[str, Any] = donnees.get("cabinet") or {}
    questions: list[dict[str, Any]] = donnees.get("questions_analytique") or []
    conclusions: list[dict[str, Any]] = (
        donnees.get("conclusions_non_verifiables") or []
    )

    exercice = mission.get("exercice")
    doc = Document()

    # En-tête cabinet (émetteur) — même style que la lettre de mission.
    doc.add_paragraph(_champ(cabinet.get("denomination")).upper())
    doc.add_paragraph(
        f"Forme juridique : {_champ(cabinet.get('forme_juridique'))} — "
        f"RCCM : {_champ(cabinet.get('rccm'))} — NCC : {_champ(cabinet.get('ncc'))}"
    )
    doc.add_paragraph(
        f"Siège : {_champ(cabinet.get('siege_social'))} — "
        f"{_champ(cabinet.get('commune'))}"
    )
    doc.add_paragraph(
        f"Centre des impôts de rattachement : {_champ(cabinet.get('centre_impots'))}"
    )

    # Destinataire
    doc.add_paragraph("")
    doc.add_paragraph(
        f"À l'attention de la Direction de {_champ(client.get('denomination'))}"
    )
    doc.add_paragraph(
        f"Siège : {_champ(client.get('siege_social'))} — "
        f"{_champ(client.get('commune'))}"
    )
    doc.add_paragraph(
        f"{_champ(cabinet.get('commune'))}, le {date.today().strftime('%d/%m/%Y')}"
    )

    # Objet
    doc.add_heading("Demande de renseignements et de documents", level=1)
    doc.add_paragraph(
        f"Objet : mission de revue fiscale — exercice {_champ(exercice)} — "
        "demande de renseignements et de documents."
    )
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph(
        "Dans le cadre de notre mission de revue fiscale de l'exercice "
        f"{_champ(exercice)}, nous vous prions de bien vouloir nous "
        "communiquer les renseignements et documents énumérés ci-après. "
        "Chaque demande est numérotée : merci de rappeler ce numéro dans "
        "votre réponse."
    )

    numero = 0

    # Section « Questions issues de la revue analytique » — omise si aucun
    # commentaire disponible.
    if questions:
        doc.add_heading("Questions issues de la revue analytique", level=2)
        doc.add_paragraph(
            "Les variations significatives relevées lors de la revue "
            "analytique N/N-1 appellent les précisions suivantes :"
        )
        for q in questions:
            numero += 1
            poste = str(q.get("poste") or "").strip() or A_COMPLETER
            gravite = str(q.get("gravite") or "").strip()
            suffixe = f" (gravité : {gravite})" if gravite else ""
            doc.add_paragraph(
                f"{numero}. Poste {poste}{suffixe} — {q.get('question')}"
            )

    # Section « Pièces et réponses attendues » — conclusions non
    # vérifiables de la dernière exécution.
    if conclusions:
        doc.add_heading("Pièces et réponses attendues", level=2)
        doc.add_paragraph(
            "Les points suivants n'ont pas pu être vérifiés faute de "
            "réponse ou de pièce justificative. Merci de fournir, pour "
            "chacun, la réponse ou la pièce manquante :"
        )
        for c in conclusions:
            numero += 1
            regle_id = str(c.get("regle_id") or "").strip() or A_COMPLETER
            libelle = str(c.get("libelle") or "").strip()
            commentaire = str(c.get("commentaire") or "").strip()
            # Intitulé/motif : libellé de la règle s'il est renseigné (et
            # pas un simple écho de l'identifiant), sinon le motif du
            # moteur (commentaire de la conclusion).
            if libelle and libelle != regle_id:
                motif = libelle
            elif commentaire:
                motif = commentaire
            else:
                motif = libelle or A_COMPLETER
            doc.add_paragraph(
                f"{numero}. [{regle_id}] — {motif} : merci de fournir la "
                "réponse ou la pièce justificative correspondante."
            )
    if numero == 0:
        doc.add_paragraph(
            "Aucune demande en attente à la date d'édition du présent "
            "document."
        )

    # Modalités de réponse
    doc.add_heading("Modalités de réponse", level=2)
    doc.add_paragraph(
        "Nous vous remercions de nous faire parvenir vos réponses et les "
        f"pièces demandées sous un délai indicatif de {DELAI_REPONSE_JOURS} "
        "jours à compter de la réception de la présente."
    )
    doc.add_paragraph(
        f"Contact du cabinet : {_champ(cabinet.get('denomination'))} — "
        f"{_champ(cabinet.get('siege_social'))} — "
        f"{_champ(cabinet.get('commune'))}. "
        f"Interlocuteur : {A_COMPLETER}."
    )
    doc.add_paragraph(
        "Nous restons à votre disposition pour toute précision et vous "
        "prions d'agréer, Madame, Monsieur, l'expression de nos "
        "salutations distinguées."
    )
    doc.add_paragraph(f"Pour le cabinet : {_champ(cabinet.get('denomination'))}")
    doc.add_paragraph("Nom et qualité : [à compléter]")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generer_demande_renseignements(
    session: Session, tenant_id: int, mission_id: int
) -> tuple[bytes, str, dict[str, int]]:
    """Contenu .docx + nom de fichier + compteurs — point d'entrée route."""
    donnees = collecter_donnees_demande(session, tenant_id, mission_id)
    contenu = rendre_demande_docx(donnees)
    nom = nom_fichier_demande(
        donnees["contribuable"].get("denomination"),
        donnees["mission"].get("exercice"),
    )
    stats = {
        "nb_questions_analytique": len(donnees["questions_analytique"]),
        "nb_pieces_attendues": len(donnees["conclusions_non_verifiables"]),
    }
    return contenu, nom, stats
