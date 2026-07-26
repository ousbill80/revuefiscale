"""Export Word du rapport de mission."""
from __future__ import annotations

import io
from collections.abc import Mapping, Sequence
from decimal import Decimal
from typing import Any

from docx import Document

from backend.restitution.passage import Passage
from backend.restitution.rapport import (
    lignes_comptes_source,
    section_fiabilite_source,
    section_note_synthese,
    section_perimetre,
    section_revue_analytique,
)
from backend.restitution.risques import ScoreRisque


def _fmt(montant: Decimal) -> str:
    return f"{montant:,.2f}".replace(",", " ").replace(".", ",")


def _ajouter_lignes_markdown_simples(doc: Document, lignes: Sequence[str]) -> None:
    """Reprend les lignes markdown du rapport texte (titres / puces / italique)."""
    for brut in lignes:
        ligne = brut.rstrip()
        if not ligne:
            continue
        if ligne.startswith("## "):
            doc.add_heading(ligne[3:].strip(), level=2)
        elif ligne.startswith("- "):
            doc.add_paragraph(ligne[2:].replace("**", "").replace("`", ""), style="List Bullet")
        elif ligne.startswith("_") and ligne.endswith("_"):
            doc.add_paragraph(ligne.strip("_"))
        else:
            doc.add_paragraph(ligne.replace("**", "").replace("`", ""))


def rendre_rapport_docx(
    *,
    meta: Mapping[str, Any],
    passage: Passage,
    conclusions: Sequence[Mapping[str, object]],
    score: ScoreRisque,
    extrait_audit: Sequence[Mapping[str, Any]],
    controles_fec: Mapping[str, Any] | None = None,
    revue_analytique: Mapping[str, Any] | None = None,
    note_synthese: Mapping[str, Any] | None = None,
) -> bytes:
    """Produit un .docx a partir des donnees deja calculees (aucun recalcul fiscal)."""
    doc = Document()
    doc.add_heading("Rapport de mission — revue fiscale", level=1)
    doc.add_heading("Identification", level=2)
    doc.add_paragraph(f"Mission : {meta.get('mission_id')}")
    doc.add_paragraph(f"Contribuable : {meta.get('contribuable_denomination') or '—'}")
    doc.add_paragraph(f"NCC : {meta.get('contribuable_ncc') or '—'}")
    doc.add_paragraph(f"Exercice : {meta.get('exercice')}")
    doc.add_paragraph(f"Version referentiel : {meta.get('version_referentiel_id')}")

    # Note de synthèse IA en tête de rapport (dernière version disponible,
    # lue en base — jamais générée à l'export). Vide → aucune section.
    _ajouter_lignes_markdown_simples(doc, section_note_synthese(note_synthese))

    _ajouter_lignes_markdown_simples(doc, section_perimetre(meta))

    # En tête de synthèse (même ordre que l'écran) : fiabilité puis revue N/N-1.
    _ajouter_lignes_markdown_simples(doc, section_fiabilite_source(controles_fec))
    _ajouter_lignes_markdown_simples(doc, section_revue_analytique(revue_analytique))

    doc.add_heading("Passage comptable / fiscal", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Regle"
    hdr[1].text = "Sens"
    hdr[2].text = "Montant (FCFA)"
    hdr[3].text = "Risque"
    for ligne in passage.lignes:
        row = table.add_row().cells
        row[0].text = str(ligne.regle_id)
        row[1].text = str(ligne.sens)
        row[2].text = _fmt(ligne.montant)
        row[3].text = str(ligne.niveau_risque)

    # Comptes à l'origine de chaque conclusion/anomalie (piste d'audit).
    for c in conclusions:
        puces = lignes_comptes_source(c)
        if not puces:
            continue
        doc.add_paragraph(
            f"Comptes à l'origine — {c.get('regle_id')} "
            f"(statut {c.get('statut') or 'anomalie'}) :"
        )
        _ajouter_lignes_markdown_simples(doc, puces)

    doc.add_paragraph(
        f"Total reintegrations : {_fmt(passage.total_reintegration)} — "
        f"deductions : {_fmt(passage.total_deduction)} — "
        f"solde net : {_fmt(passage.solde_net)}"
    )
    doc.add_paragraph(
        f"Score risque (heuristique, non CGI) : {score.score} — {score.comptages}"
    )
    if score.avertissement:
        doc.add_paragraph(str(score.avertissement))

    doc.add_heading("Journal d audit (extrait)", level=2)
    for entree in extrait_audit:
        doc.add_paragraph(
            f"{entree.get('horodatage')} | {entree.get('acteur')} | {entree.get('action')}"
        )

    doc.add_paragraph(
        "Note : montants issus du referentiel YAML (mentions a confirmer possibles)."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
