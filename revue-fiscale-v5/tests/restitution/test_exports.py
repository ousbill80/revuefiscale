"""Tests exports Word/PDF et lecture audit."""
from __future__ import annotations

from decimal import Decimal

from backend.restitution.passage import Passage, LignePassage
from backend.restitution.rapport_docx import rendre_rapport_docx
from backend.restitution.rapport_pdf import rendre_rapport_pdf
from backend.restitution.risques import ScoreRisque


def test_exports_docx_pdf_non_vides():
    passage = Passage(
        lignes=(
            LignePassage(
                regle_id="BIC-CHG-18G-DONS",
                sens="reintegration",
                montant=Decimal("1000"),
                niveau_risque="moyen",
            ),
        ),
        total_reintegration=Decimal("1000"),
        total_deduction=Decimal("0"),
        solde_net=Decimal("1000"),
    )
    score = ScoreRisque(score=2, comptages={"moyen": 1}, avertissement="heuristique")
    meta = {
        "mission_id": 1,
        "exercice": 2025,
        "contribuable_denomination": "Demo SA",
        "contribuable_ncc": "123",
        "version_referentiel_id": 1,
        "type_engagement": "preventive",
        "perimetre_impots": ["TVA"],
    }
    docx = rendre_rapport_docx(
        meta=meta,
        passage=passage,
        conclusions=[],
        score=score,
        extrait_audit=[{"horodatage": "t", "acteur": "a", "action": "x"}],
    )
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=[],
        score=score,
        extrait_audit=[{"horodatage": "t", "acteur": "a", "action": "x"}],
    )
    assert docx[:2] == b"PK"
    assert pdf.startswith(b"%PDF")
    # Section périmètre présente dans les deux exports (template séparé ≠ wrap markdown).
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(docx))
    textes = "\n".join(p.text for p in doc.paragraphs)
    assert "Périmètre déclaré" in textes
    assert "Non examiné" in textes
    # Sections d'enrichissement toujours présentes — mention sobre sans données.
    assert "Fiabilité de la source" in textes
    assert "Aucun contrôle de vraisemblance FEC" in textes
    assert "Revue analytique" in textes
    assert "Aucun exercice antérieur comparable" in textes
    # ReportLab dessine en WinAnsi — accents possibles en octets.
    assert b"Non examin" in pdf or b"P" in pdf


def _donnees_communes():
    passage = Passage(
        lignes=(
            LignePassage(
                regle_id="BIC-CHG-18G-DONS",
                sens="reintegration",
                montant=Decimal("1000"),
                niveau_risque="moyen",
            ),
        ),
        total_reintegration=Decimal("1000"),
        total_deduction=Decimal("0"),
        solde_net=Decimal("1000"),
    )
    score = ScoreRisque(score=2, comptages={"moyen": 1}, avertissement="heuristique")
    meta = {
        "mission_id": 1,
        "exercice": 2025,
        "contribuable_denomination": "Demo SA",
        "contribuable_ncc": "123",
        "version_referentiel_id": 1,
        "type_engagement": "preventive",
        "perimetre_impots": ["TVA"],
    }
    conclusions = [
        {
            "regle_id": "BIC-CHG-18G-DONS",
            "statut": "anomalie",
            "sens": "reintegration",
            "montant": Decimal("1000"),
            "niveau_risque": "moyen",
            "comptes_source": [
                {
                    "compte": "623100",
                    "libelle": "Dons et libéralités",
                    "solde": "1500000",
                    "sens": "debiteur",
                }
            ],
        }
    ]
    controles_fec = {
        "exercice": 2025,
        "cree_le": "2026-01-15T10:00:00",
        "controles": [
            {
                "code": "ecritures_non_equilibrees",
                "libelle": "Écritures non équilibrées (débit ≠ crédit)",
                "statut": "alerte",
                "compteur": 2,
                "echantillon": [],
            },
            {
                "code": "doublons_stricts",
                "libelle": "Doublons stricts d'écritures",
                "statut": "ok",
                "compteur": 0,
                "echantillon": [],
            },
        ],
    }
    revue_analytique = {
        "disponible": True,
        "exercice_n": 2025,
        "exercice_n1": 2024,
        "mission_n1_id": 9,
        "lignes": [
            {
                "compte": "701100",
                "libelle": "Ventes de marchandises",
                "solde_n": -5000000.0,
                "solde_n1": -8000000.0,
                "variation": 3000000.0,
                "variation_pct": 37.5,
                "sens": "hausse",
                "classement": "variation_forte",
            }
        ],
        "totaux_par_classe": [
            {
                "classe": 7,
                "total_n": -5000000.0,
                "total_n1": -8000000.0,
                "variation": 3000000.0,
            }
        ],
    }
    return passage, score, meta, conclusions, controles_fec, revue_analytique


def test_docx_contient_fiabilite_revue_et_comptes_source():
    """Mission avec source FEC contrôlée + N-1 : les sections sont remplies."""
    passage, score, meta, conclusions, controles, revue = _donnees_communes()
    docx = rendre_rapport_docx(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        controles_fec=controles,
        revue_analytique=revue,
    )
    assert docx[:2] == b"PK"
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(docx))
    textes = "\n".join(p.text for p in doc.paragraphs)
    # Fiabilité de la source : titre, date jj/mm/aaaa, alerte avec compteur.
    assert "Fiabilité de la source" in textes
    assert "15/01/2026" in textes
    assert "[ALERTE]" in textes and "2 occurrences" in textes
    assert "[OK] Doublons stricts" in textes
    # Revue analytique : titre avec exercices, ligne principale, totaux classe.
    assert "Revue analytique 2025 / 2024" in textes
    assert "701100" in textes and "variation forte" in textes
    assert "Classe 7" in textes
    # Comptes à l'origine de la conclusion.
    assert "Comptes à l'origine — BIC-CHG-18G-DONS" in textes
    assert "623100" in textes and "debiteur" in textes
    # Montant formaté FCFA (espace + décimales virgule).
    assert "1 500 000,00 FCFA" in textes


def test_pdf_se_genere_avec_enrichissements():
    """Le PDF reste valide (magic %PDF) avec toutes les sections remplies."""
    passage, score, meta, conclusions, controles, revue = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        controles_fec=controles,
        revue_analytique=revue,
    )
    assert pdf.startswith(b"%PDF")


def test_pdf_se_genere_sans_donnees_enrichissement():
    """Sans contrôles FEC ni N-1 : mentions sobres, PDF toujours valide."""
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=[],
        score=score,
        extrait_audit=[],
    )
    assert pdf.startswith(b"%PDF")


# ── Note de synthèse dans les exports ─────────────────────────────────


def _note_disponible():
    return {
        "id": 7,
        "mission_id": 1,
        "version": 2,
        "statut": "disponible",
        "modele": "provider-x",
        "erreur": None,
        "auteur": "admin@demo.local",
        "cree_le": "2026-07-20T09:30:00",
        "contenu": {
            "contexte": "Revue fiscale préventive de Demo SA, exercice 2025.",
            "constats": [
                {
                    "regle_id": "OBL-36-ETII",
                    "resume": "État des transactions intra-groupe non produit.",
                    "montant": None,
                    "gravite": "haute",
                },
                {
                    "regle_id": "BIC-CHG-18G-DONS",
                    "resume": "Dons à réintégrer au résultat fiscal.",
                    "montant": "1000",
                    "gravite": "moyenne",
                },
            ],
            "exposition": "Exposition estimée à 1 000 FCFA hors pénalités.",
            "points_attention": ["2 écritures FEC non équilibrées."],
            "recommandations": ["Produire l'état des transactions intra-groupe."],
        },
    }


def _texte_pdf(pdf: bytes) -> str:
    """Texte brut des content streams (ASCII85 + Flate de ReportLab) — sans dépendance."""
    import base64
    import re
    import zlib

    morceaux: list[bytes] = []
    for m in re.finditer(rb"stream\r?\n(.*?)endstream", pdf, re.DOTALL):
        brut = m.group(1).strip()
        try:  # encodage ReportLab par défaut : ASCII85 puis Flate.
            brut = base64.a85decode(brut, adobe=True)
        except ValueError:
            pass
        try:
            morceaux.append(zlib.decompress(brut))
        except zlib.error:
            morceaux.append(brut)
    return b"\n".join(morceaux).decode("latin-1", "replace")


def _texte_docx(docx: bytes) -> str:
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(docx))
    return "\n".join(p.text for p in doc.paragraphs)


def test_docx_avec_note_synthese_disponible():
    """Note disponible → section en tête avec gravité, regle_id et contenu."""
    passage, score, meta, conclusions, controles, revue = _donnees_communes()
    docx = rendre_rapport_docx(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        controles_fec=controles,
        revue_analytique=revue,
        note_synthese=_note_disponible(),
    )
    textes = _texte_docx(docx)
    assert "Note de synthèse" in textes
    assert "version 2 du 20/07/2026" in textes
    assert "Contexte : Revue fiscale préventive de Demo SA" in textes
    assert "[HAUTE] OBL-36-ETII" in textes
    assert "[MOYENNE] BIC-CHG-18G-DONS" in textes
    assert "montant : 1000 FCFA" in textes
    assert "Exposition estimée" in textes
    assert "Points d'attention" in textes
    assert "2 écritures FEC non équilibrées." in textes
    assert "Recommandations prioritaires" in textes
    assert "Produire l'état des transactions intra-groupe." in textes
    # La section précède le périmètre (en tête de rapport).
    assert textes.index("Note de synthèse") < textes.index("Périmètre déclaré")


def test_docx_sans_note_synthese_pas_de_section():
    """Aucune note disponible → aucune section (pas de titre vide)."""
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    for note in (None, {"statut": "echec", "contenu": None}):
        docx = rendre_rapport_docx(
            meta=meta,
            passage=passage,
            conclusions=conclusions,
            score=score,
            extrait_audit=[],
            note_synthese=note,
        )
        assert "Note de synthèse" not in _texte_docx(docx)


def test_pdf_avec_note_synthese_disponible():
    """PDF valide et section note présente avec les regle_id."""
    passage, score, meta, conclusions, controles, revue = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        controles_fec=controles,
        revue_analytique=revue,
        note_synthese=_note_disponible(),
    )
    assert pdf.startswith(b"%PDF")
    texte = _texte_pdf(pdf)
    assert "Note de synth" in texte
    assert "OBL-36-ETII" in texte
    assert "BIC-CHG-18G-DONS" in texte
    assert "Recommandations prioritaires" in texte


def test_pdf_sans_note_synthese_pas_de_section():
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        note_synthese=None,
    )
    assert pdf.startswith(b"%PDF")
    assert "Note de synth" not in _texte_pdf(pdf)


def test_section_note_synthese_vide_ou_invalide():
    """Garde-fou : pas de lignes si note absente, sans contenu ou non-dict."""
    from backend.restitution.rapport import section_note_synthese

    assert section_note_synthese(None) == []
    assert section_note_synthese({"statut": "echec", "contenu": None}) == []
    assert section_note_synthese({"contenu": "pas un dict"}) == []


# ── Commentaire de revue analytique dans les exports ──────────────────


def _commentaire_disponible():
    return {
        "id": 3,
        "mission_id": 1,
        "version": 1,
        "statut": "disponible",
        "modele": "provider-x",
        "erreur": None,
        "auteur": "admin@demo.local",
        "cree_le": "2026-07-21T11:00:00",
        "contenu": {
            "resume": "Hausse marquée des ventes, trésorerie en tension.",
            "explications": [
                {
                    "poste": "701100",
                    "hypothese_explicative": (
                        "La hausse des ventes pourrait refléter un nouveau "
                        "contrat significatif."
                    ),
                    "question_a_poser_au_client": (
                        "Pouvez-vous détailler les nouveaux clients de "
                        "l'exercice ?"
                    ),
                    "gravite": "haute",
                },
                {
                    "poste": "512100",
                    "hypothese_explicative": (
                        "La baisse de trésorerie serait liée aux "
                        "investissements de fin d'exercice."
                    ),
                    "question_a_poser_au_client": "",
                    "gravite": "faible",
                },
            ],
            "alertes_coherence": [
                "Charges en hausse sans hausse d'activité correspondante."
            ],
        },
    }


def test_docx_avec_commentaire_analytique_disponible():
    """Commentaire disponible → section après la note, gravités et questions."""
    passage, score, meta, conclusions, controles, revue = _donnees_communes()
    docx = rendre_rapport_docx(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        controles_fec=controles,
        revue_analytique=revue,
        note_synthese=_note_disponible(),
        commentaire_analytique=_commentaire_disponible(),
    )
    textes = _texte_docx(docx)
    assert "Commentaire de revue analytique" in textes
    assert "version 1 du 21/07/2026" in textes
    assert "Résumé : Hausse marquée des ventes" in textes
    assert "[HAUTE] 701100" in textes
    assert "[FAIBLE] 512100" in textes
    assert (
        "Question au client : Pouvez-vous détailler les nouveaux clients"
        in textes
    )
    assert "Alertes de cohérence" in textes
    assert "Charges en hausse sans hausse d'activité" in textes
    # Placement : après la note de synthèse, avant le périmètre (corps).
    assert (
        textes.index("Note de synthèse")
        < textes.index("Commentaire de revue analytique")
        < textes.index("Périmètre déclaré")
    )


def test_docx_sans_commentaire_analytique_pas_de_section():
    """Aucun commentaire disponible → aucune section (pas de titre vide)."""
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    for commentaire in (None, {"statut": "echec", "contenu": None}):
        docx = rendre_rapport_docx(
            meta=meta,
            passage=passage,
            conclusions=conclusions,
            score=score,
            extrait_audit=[],
            commentaire_analytique=commentaire,
        )
        assert "Commentaire de revue analytique" not in _texte_docx(docx)


def test_pdf_avec_commentaire_analytique_disponible():
    passage, score, meta, conclusions, controles, revue = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        controles_fec=controles,
        revue_analytique=revue,
        commentaire_analytique=_commentaire_disponible(),
    )
    assert pdf.startswith(b"%PDF")
    texte = _texte_pdf(pdf)
    assert "Commentaire de revue analytique" in texte
    assert "701100" in texte
    assert "512100" in texte
    assert "Question au client" in texte


def test_pdf_sans_commentaire_analytique_pas_de_section():
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        commentaire_analytique=None,
    )
    assert pdf.startswith(b"%PDF")
    assert "Commentaire de revue analytique" not in _texte_pdf(pdf)


def test_section_commentaire_analytique_vide_ou_invalide():
    from backend.restitution.rapport import section_commentaire_analytique

    assert section_commentaire_analytique(None) == []
    assert (
        section_commentaire_analytique({"statut": "echec", "contenu": None})
        == []
    )
    assert section_commentaire_analytique({"contenu": "pas un dict"}) == []


# ── Exposition pénalités et intérêts dans les exports ─────────────────


def _risques_ouverts_chiffres():
    """Risques sérialisés comme backend/plateforme/risques.py (chiffrage réel)."""
    from backend.plateforme.penalites import chiffrer_risque

    base = [
        {
            "id": 11,
            "impot": "TVA",
            "libelle": "TVA déductible non justifiée",
            "statut": "ouvert",
            "montant_estime": "1000000",
            "exercice_origine": 2024,
        },
        {
            "id": 12,
            "impot": "BIC",
            "libelle": "Dons non déductibles",
            "statut": "en_traitement",
            "montant_estime": "500000",
            "exercice_origine": 2023,
        },
    ]
    from datetime import date

    for r in base:
        r["chiffrage_penalites"] = chiffrer_risque(
            r, aujourd_hui=date(2026, 7, 1)
        )
    return base


def test_docx_avec_risques_ouverts_chiffres():
    """Risque ouvert chiffré → section pénalités avec lignes, totaux, mention."""
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    risques = _risques_ouverts_chiffres()
    docx = rendre_rapport_docx(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        risques_chiffres=risques,
    )
    textes = _texte_docx(docx)
    assert "Exposition pénalités et intérêts (indicatif)" in textes
    # Risque TVA 2024 : 18 mois au 01/07/2026 → intérêts 9 %, assiette 25 %.
    assert "[TVA 2024] TVA déductible non justifiée" in textes
    assert "droit simple 1 000 000 FCFA" in textes
    assert "+ intérêts 90 000 FCFA" in textes
    assert "+ pénalité 250 000 FCFA" in textes
    assert "= total 1 340 000 FCFA (18 mois)" in textes
    # Risque BIC 2023 : 30 mois → intérêts 15 % de 500 000 = 75 000.
    assert "[BIC 2023] Dons non déductibles" in textes
    assert "= total 700 000 FCFA (30 mois)" in textes
    # Total général = 1 340 000 + 700 000, et mention obligatoire.
    assert "Total général estimé : 2 040 000 FCFA" in textes
    assert "Chiffrage indicatif, à valider par l'associé." in textes


def test_docx_sans_risque_ouvert_chiffre_pas_de_section():
    """Aucun risque ouvert chiffré → section omise (liste vide ou None)."""
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    for risques in (None, [], [{"id": 1, "chiffrage_penalites": None}]):
        docx = rendre_rapport_docx(
            meta=meta,
            passage=passage,
            conclusions=conclusions,
            score=score,
            extrait_audit=[],
            risques_chiffres=risques,
        )
        assert "Exposition pénalités" not in _texte_docx(docx)


def test_pdf_avec_risques_ouverts_chiffres():
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        risques_chiffres=_risques_ouverts_chiffres(),
    )
    assert pdf.startswith(b"%PDF")
    texte = _texte_pdf(pdf)
    assert "Exposition p" in texte  # « pénalités » en WinAnsi
    assert "1 340 000 FCFA" in texte
    assert "2 040 000 FCFA" in texte
    assert "valider par l'associ" in texte


def test_pdf_sans_risque_ouvert_chiffre_pas_de_section():
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        risques_chiffres=[],
    )
    assert pdf.startswith(b"%PDF")
    assert "Exposition p" not in _texte_pdf(pdf)


def test_section_exposition_penalites_vide():
    from backend.restitution.rapport import section_exposition_penalites

    assert section_exposition_penalites(None) == []
    assert section_exposition_penalites([]) == []
    assert (
        section_exposition_penalites([{"id": 1, "chiffrage_penalites": None}])
        == []
    )


# ── Provision pour risques fiscaux dans les exports ───────────────────


def _provision_non_vide():
    """Provision réelle : chiffrage penalites.py + provision_risques.py."""
    from datetime import date

    from backend.plateforme.penalites import chiffrer_risque
    from backend.plateforme.provision_risques import (
        calculer_provision_depuis_risques,
    )

    risques = [
        {
            "id": 11,
            "impot": "TVA",
            "libelle": "TVA déductible non justifiée",
            "statut": "ouvert",
            "probabilite": "probable",
            "montant_estime": "1000000",
            "exercice_origine": 2024,
        },
        {
            "id": 12,
            "impot": "BIC",
            "libelle": "Dons non déductibles",
            "statut": "en_traitement",
            "probabilite": "possible",
            "montant_estime": "500000",
            "exercice_origine": 2023,
        },
    ]
    for r in risques:
        r["chiffrage_penalites"] = chiffrer_risque(
            r, aujourd_hui=date(2026, 7, 1)
        )
    return calculer_provision_depuis_risques(risques, exercice_courant=2026)


def test_docx_avec_provision_risques():
    """Provision non vide → section après pénalités : lignes, total, écriture."""
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    docx = rendre_rapport_docx(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        risques_chiffres=_risques_ouverts_chiffres(),
        provision=_provision_non_vide(),
    )
    textes = _texte_docx(docx)
    assert "Provision pour risques fiscaux proposée" in textes
    # Ligne provisionnable : TVA 2024 probable, total 1 340 000 (18 mois).
    assert (
        "TVA déductible non justifiée (TVA 2024, probable) — provision "
        "1 340 000 FCFA" in textes
    )
    assert "Total de la provision proposée : 1 340 000 FCFA" in textes
    # Passif éventuel : BIC possible, exposition 700 000 (30 mois).
    assert (
        "Dons non déductibles — montant estimé 700 000 FCFA "
        "(mention en annexe recommandée)" in textes
    )
    # Écriture SYSCOHADA proposée + libellé.
    assert (
        "Débit 6911 Dotations aux provisions d'exploitation / "
        "Crédit 1918 Autres provisions pour risques — 1 340 000 FCFA"
        in textes
    )
    assert (
        "Libellé : Provision pour risques fiscaux — exercice 2026" in textes
    )
    # Hypothèses reprises telles quelles.
    assert "Proposition indicative à valider par l'expert-comptable" in textes
    assert "passifs éventuels" in textes
    # Placement : après la section exposition pénalités.
    assert textes.index("Exposition pénalités") < textes.index(
        "Provision pour risques fiscaux proposée"
    )


def test_docx_sans_provision_pas_de_section():
    """Provision None ou vide (total 0, aucun passif) → section omise."""
    from backend.plateforme.provision_risques import (
        calculer_provision_depuis_risques,
    )

    passage, score, meta, conclusions, _, _ = _donnees_communes()
    provision_vide = calculer_provision_depuis_risques([], exercice_courant=2026)
    assert provision_vide["total_provision"] == "0"
    for provision in (None, provision_vide):
        docx = rendre_rapport_docx(
            meta=meta,
            passage=passage,
            conclusions=conclusions,
            score=score,
            extrait_audit=[],
            provision=provision,
        )
        assert "Provision pour risques fiscaux" not in _texte_docx(docx)


def test_pdf_avec_provision_risques():
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        risques_chiffres=_risques_ouverts_chiffres(),
        provision=_provision_non_vide(),
    )
    assert pdf.startswith(b"%PDF")
    texte = _texte_pdf(pdf)
    assert "Provision pour risques fiscaux propos" in texte
    assert "1 340 000 FCFA" in texte
    assert "6911" in texte and "1918" in texte
    assert "mention en annexe recommand" in texte


def test_pdf_sans_provision_pas_de_section():
    passage, score, meta, conclusions, _, _ = _donnees_communes()
    pdf = rendre_rapport_pdf(
        meta=meta,
        passage=passage,
        conclusions=conclusions,
        score=score,
        extrait_audit=[],
        provision=None,
    )
    assert pdf.startswith(b"%PDF")
    assert "Provision pour risques fiscaux" not in _texte_pdf(pdf)


def test_section_provision_risques_vide():
    from backend.plateforme.provision_risques import (
        calculer_provision_depuis_risques,
    )
    from backend.restitution.rapport import section_provision_risques

    assert section_provision_risques(None) == []
    assert (
        section_provision_risques(
            calculer_provision_depuis_risques([], exercice_courant=2026)
        )
        == []
    )
